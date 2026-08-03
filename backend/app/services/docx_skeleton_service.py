"""Fill a cloned DOCX company template with composed candidate content."""

from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Any

from app.core.logging import logger
from app.models.format_schema import normalize_format_metadata, to_heading_title_case
from app.services.doc_converter import DocConversionError, convert_doc_to_docx
from app.services.docx_layout import (
    add_bullet_paragraph,
    add_left_right_line,
    split_left_right_meta,
    summary_to_bullets,
    usable_content_width,
)


def render_from_docx_skeleton(
    template_path: Path,
    document: dict[str, Any],
    metadata: dict[str, Any] | None,
    *,
    source_type: str = "docx",
) -> bytes | None:
    """
    Clone a DOCX (or converted DOC) skeleton: keep headers/footers/styles,
    clear body sample content, inject composed resume content.
    Returns None to signal caller should fall back to schema renderer.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from io import BytesIO
    except ImportError:
        return None

    path = Path(template_path)
    if not path.exists():
        return None

    work_path = path
    tmp_converted: Path | None = None
    if source_type == "doc" or path.suffix.lower() == ".doc":
        try:
            tmp_converted = convert_doc_to_docx(path, output_dir=path.parent)
            work_path = tmp_converted
        except DocConversionError as exc:
            logger.warning("skeleton_doc_convert_failed", error=str(exc))
            return None

    if work_path.suffix.lower() != ".docx":
        return None

    metadata = normalize_format_metadata(metadata)
    styling = metadata.get("styling") or {}
    # Client policy: Arial 10 for name, headlines, and body.
    font_family = str(styling.get("font_family") or "Arial")
    body_size = float(styling.get("font_size_body") or 10)
    header_size = float(styling.get("font_size_header") or body_size or 10)
    name_size = float(styling.get("font_size_name") or body_size or 10)
    color_text = _rgb(styling.get("color_text"), (0, 0, 0))
    color_muted = _rgb(styling.get("color_muted"), (0x33, 0x33, 0x33))
    space_after = float(styling.get("space_after_para") or 6)

    try:
        doc = Document(str(work_path))
    except Exception as exc:
        logger.warning("skeleton_open_failed", error=str(exc))
        return None

    # Clear body paragraphs/tables but keep section properties + headers/footers.
    _clear_body(doc)

    try:
        normal = doc.styles["Normal"]
        normal.font.name = font_family
        normal.font.size = Pt(body_size)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)
    except Exception:
        pass

    usable_width = usable_content_width(doc)

    header = document.get("header") or {}
    name = str(header.get("name") or "Candidate").strip() or "Candidate"
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(name)
    name_run.bold = True
    name_run.font.name = font_family
    name_run.font.size = Pt(name_size)
    name_run.font.color.rgb = color_text
    name_run.italic = False
    name_run.underline = False

    # Client policy: name only — never email/phone/address on generated resumes.

    for section in document.get("sections") or []:
        if not isinstance(section, dict):
            continue
        title = str(section.get("title") or "").strip()
        if title:
            title_para = doc.add_paragraph()
            display_title = to_heading_title_case(title, keep_colon=False)
            title_run = title_para.add_run(display_title)
            title_run.bold = True
            title_run.italic = False
            title_run.underline = False
            title_run.font.name = font_family
            title_run.font.size = Pt(header_size)
            title_run.font.color.rgb = color_text
            title_para.paragraph_format.space_before = Pt(12)
            title_para.paragraph_format.space_after = Pt(space_after)

        content = section.get("content")
        stype = str(section.get("type") or "").lower()
        if not content:
            continue
        if stype == "text":
            # Format templates use bullet summary — never dump as one paragraph.
            bullets = summary_to_bullets(content)
            if bullets:
                for text in bullets:
                    _add_bullet(doc, text, font_family, body_size, color_text)
            else:
                _add_plain_or_split_line(
                    doc, str(content), font_family, body_size, color_text, color_muted, usable_width
                )
        elif stype == "skills":
            _add_skills(doc, content, font_family, body_size, color_text)
        elif stype in {"experience", "projects", "education"}:
            for item in content if isinstance(content, list) else [content]:
                _add_record(
                    doc,
                    item,
                    font_family,
                    body_size,
                    color_text,
                    color_muted,
                    space_after,
                    usable_width,
                )
        else:
            items = content if isinstance(content, list) else [content]
            for item in items:
                if isinstance(item, dict):
                    _add_record(
                        doc,
                        item,
                        font_family,
                        body_size,
                        color_text,
                        color_muted,
                        space_after,
                        usable_width,
                    )
                    continue
                text = str(item or "").strip()
                if not text:
                    continue
                left, right = split_left_right_meta(text)
                if right:
                    add_left_right_line(
                        doc,
                        left,
                        right,
                        font_family=font_family,
                        body_size=body_size,
                        color_text=color_text,
                        color_muted=color_muted,
                        usable_width=usable_width,
                        bold_left=False,
                        space_before=2,
                        space_after=2,
                    )
                else:
                    _add_bullet(doc, text, font_family, body_size, color_text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    result = buffer.read()

    if tmp_converted and tmp_converted.exists() and tmp_converted != path:
        try:
            if "tmp" in str(tmp_converted).lower() or tempfile.gettempdir() in str(tmp_converted):
                tmp_converted.unlink(missing_ok=True)
        except Exception:
            pass

    return result if result else None


def _rgb(value: Any, default: tuple[int, int, int]) -> Any:
    from docx.shared import RGBColor
    from app.models.format_schema import parse_hex_rgb

    try:
        return RGBColor(*parse_hex_rgb(str(value or "")))
    except Exception:
        return RGBColor(*default)


def _clear_body(doc: Any) -> None:
    from docx.oxml.ns import qn

    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _add_bullet(doc: Any, text: str, font_family: str, body_size: float, color: Any) -> None:
    add_bullet_paragraph(
        doc,
        text,
        font_family=font_family,
        body_size=body_size,
        color=color,
        space_after=2,
    )


def _add_plain_or_split_line(
    doc: Any,
    text: str,
    font_family: str,
    body_size: float,
    color_text: Any,
    color_muted: Any,
    usable_width: Any,
) -> None:
    from docx.shared import Pt

    left, right = split_left_right_meta(text)
    if right:
        add_left_right_line(
            doc,
            left,
            right,
            font_family=font_family,
            body_size=body_size,
            color_text=color_text,
            color_muted=color_muted,
            usable_width=usable_width,
            bold_left=False,
            space_before=0,
            space_after=4,
        )
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_family
    run.font.size = Pt(body_size)
    run.font.color.rgb = color_text
    run.italic = False
    run.underline = False


def _add_skills(doc: Any, content: Any, font_family: str, body_size: float, color: Any) -> None:
    from docx.shared import Pt

    if isinstance(content, dict):
        for category, skill_list in content.items():
            skills = [str(s).strip() for s in (skill_list or []) if str(s).strip()]
            if not skills:
                continue
            p = doc.add_paragraph()
            cat = p.add_run(f"{category}: ")
            cat.bold = True
            cat.italic = False
            cat.underline = False
            cat.font.name = font_family
            cat.font.size = Pt(body_size)
            cat.font.color.rgb = color
            rest = p.add_run(", ".join(skills))
            rest.font.name = font_family
            rest.font.size = Pt(body_size)
            rest.font.color.rgb = color
            rest.italic = False
            rest.underline = False
        return
    skills = [str(s).strip() for s in (content if isinstance(content, list) else [content]) if str(s).strip()]
    if skills:
        p = doc.add_paragraph()
        run = p.add_run(", ".join(skills))
        run.font.name = font_family
        run.font.size = Pt(body_size)
        run.font.color.rgb = color
        run.italic = False
        run.underline = False


def _add_record(
    doc: Any,
    item: Any,
    font_family: str,
    body_size: float,
    color_text: Any,
    color_muted: Any,
    space_after: float,
    usable_width: Any,
) -> None:
    from docx.shared import Pt

    if not isinstance(item, dict):
        text = str(item or "").strip()
        if text:
            left, right = split_left_right_meta(text)
            if right:
                add_left_right_line(
                    doc,
                    left,
                    right,
                    font_family=font_family,
                    body_size=body_size,
                    color_text=color_text,
                    color_muted=color_muted,
                    usable_width=usable_width,
                    bold_left=True,
                )
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = font_family
                run.font.size = Pt(body_size)
                run.font.color.rgb = color_text
                run.italic = False
                run.underline = False
        return

    title = str(
        item.get("title")
        or item.get("role")
        or item.get("degree")
        or item.get("name")
        or ""
    ).strip()
    company = str(
        item.get("company")
        or item.get("institution")
        or item.get("organization")
        or ""
    ).strip()
    location = str(item.get("location") or "").strip()
    dates = str(
        item.get("duration") or item.get("dates") or item.get("year") or item.get("link") or item.get("url") or ""
    ).strip()
    project_name = str(item.get("project") or item.get("project_name") or "").strip()
    bullets = item.get("description") or item.get("bullets") or item.get("achievements") or []

    # Prefer explicit date/url fields; else peel from company/title strings.
    primary = company or title
    if primary and not dates:
        primary, dates = split_left_right_meta(primary)
        if company:
            company = primary
        else:
            title = primary

    left_label = primary
    if company and location and location.lower() not in company.lower():
        left_label = f"{company}, {location}" if company else primary

    if left_label or dates:
        add_left_right_line(
            doc,
            left_label or title,
            dates,
            font_family=font_family,
            body_size=body_size,
            color_text=color_text,
            color_muted=color_muted,
            usable_width=usable_width,
            bold_left=True,
            space_before=6,
            space_after=0,
        )

    if company and title and company != title:
        role = doc.add_paragraph()
        role.paragraph_format.space_after = Pt(2)
        role_run = role.add_run(title)
        role_run.bold = True
        role_run.italic = False
        role_run.underline = False
        role_run.font.name = font_family
        role_run.font.size = Pt(body_size)
        role_run.font.color.rgb = color_text

    if project_name:
        proj_left, proj_right = split_left_right_meta(project_name)
        if proj_right:
            add_left_right_line(
                doc,
                proj_left,
                proj_right,
                font_family=font_family,
                body_size=body_size,
                color_text=color_text,
                color_muted=color_muted,
                usable_width=usable_width,
                bold_left=True,
                space_before=0,
                space_after=2,
            )
        else:
            proj = doc.add_paragraph()
            proj.paragraph_format.space_after = Pt(2)
            proj_run = proj.add_run(project_name)
            proj_run.bold = True
            proj_run.italic = False
            proj_run.underline = False
            proj_run.font.name = font_family
            proj_run.font.size = Pt(body_size)
            proj_run.font.color.rgb = color_text

    if isinstance(bullets, str):
        bullets = [bullets]
    for bullet in bullets or []:
        text = str(bullet or "").strip()
        if not text:
            continue
        left, right = split_left_right_meta(text)
        if right and len(left.split()) <= 12:
            # Short label + date/link → right-align meta (rare in bullets).
            add_left_right_line(
                doc,
                left,
                right,
                font_family=font_family,
                body_size=body_size,
                color_text=color_text,
                color_muted=color_muted,
                usable_width=usable_width,
                bold_left=False,
                space_before=1,
                space_after=2,
            )
        else:
            _add_bullet(doc, text, font_family, body_size, color_text)
