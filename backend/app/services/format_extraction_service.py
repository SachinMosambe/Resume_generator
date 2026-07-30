import re
import tempfile
import logging
from pathlib import Path
from typing import Any

try:
    from langsmith import traceable
except ImportError:  # pragma: no cover
    def traceable(*args, **kwargs):
        def decorator(func):
            return func
        return decorator

logger = logging.getLogger(__name__)


class FormatExtractionError(ValueError):
    pass


SECTION_ALIASES: dict[str, list[str]] = {
    "header": ["header", "contact", "profile"],
    "summary": ["summary", "professional summary", "profile summary", "objective"],
    "experience": ["experience", "work experience", "professional experience", "employment"],
    "education": ["education", "academic", "academics", "qualifications"],
    "skills": ["skills", "technical skills", "core skills", "competencies"],
    "projects": ["projects", "selected projects", "project experience"],
    "certifications": ["certifications", "certificates", "licenses"],
}

DEFAULT_SECTIONS = [
    "header",
    "summary",
    "skills",
    "experience",
    "projects",
    "education",
    "certifications",
    "achievements",
]


class FormatExtractionService:
    allowed_extensions = {".pdf", ".docx", ".doc"}

    @traceable(run_type="retriever", tags=["format", "extraction", "client-template"])
    def extract(self, filename: str, content: bytes) -> dict[str, Any]:
        logger.info(f"Extracting client format from: {filename} ({len(content)} bytes)")
        
        ext = Path(filename or "").suffix.lower()
        if ext not in self.allowed_extensions:
            raise FormatExtractionError("Client format must be a PDF, DOC, or DOCX file")
        if not content:
            raise FormatExtractionError("Client format file is empty")

        text, styling, logos, header_text, footer_text = self._extract_document(ext, content)
        sections, section_labels = self._infer_sections_with_labels(text)

        # Prefer formal header/footer branding; fall back to body heuristics.
        header_sign = self._company_sign_from_text(header_text)
        footer_sign = self._company_sign_from_text(footer_text)
        body_sign = self._infer_company_header(text)
        company_header = header_sign or body_sign
        company_footer = footer_sign or header_sign or body_sign

        # If branding text mentions Aptino but no logo was extracted, inject built-in Aptino logo.
        logos = list(logos or [])
        branding_text = " ".join(
            [
                header_text or "",
                footer_text or "",
                " ".join((company_footer or {}).get("lines") or []) if isinstance(company_footer, dict) else "",
                " ".join((company_header or {}).get("lines") or []) if isinstance(company_header, dict) else "",
            ]
        ).lower()
        if not logos and "aptino" in branding_text:
            try:
                from app.services.aptino_template import get_aptino_default_metadata

                fallback_logos = get_aptino_default_metadata().get("logos") or []
                if fallback_logos:
                    logos = list(fallback_logos)
                    logger.info("Injected built-in Aptino logo because branding text matched Aptino")
            except Exception as exc:
                logger.debug("Aptino logo fallback failed: %s", exc)
        if not company_footer and "aptino" in branding_text:
            try:
                from app.services.aptino_template import get_aptino_default_metadata

                company_footer = get_aptino_default_metadata().get("company_footer")
                company_header = company_header or company_footer
            except Exception:
                pass

        preview_text = self._build_format_preview(text, sections, section_labels)

        logger.info(
            "Extraction complete: %s sections, %s logos, header_branding=%s, footer_branding=%s",
            len(sections),
            len(logos),
            bool(company_header),
            bool(company_footer),
        )
        if logos:
            for i, logo in enumerate(logos):
                data_len = len(logo.get("data", "")) if logo.get("data") else 0
                logger.info("  Logo %s: %s (%s chars, source=%s)", i + 1, logo.get("position"), data_len, logo.get("source"))

        return {
            "source_filename": filename,
            "source_type": ext.replace(".", ""),
            "sections": sections,
            "section_order": list(range(len(sections))),
            "styling": styling,
            "field_mapping": self._build_field_mapping(sections, section_labels),
            "section_labels": section_labels,
            "layout": {
                "type": self._infer_layout(text),
                "logo_position": "top_right",
                "name_position": "top_left",
                "company_header": "center" if company_header else None,
                "company_footer": "center" if company_footer else None,
                "dates": "right_aligned",
                "section_dividers": True,
            },
            "company_header": company_header,
            "company_footer": company_footer,
            "preview_text": preview_text,
            "logos": logos,
            "logo_count": len(logos),
            # Intentionally blank: sample resume candidate names must not be copied.
            "header_text": "",
            "footer_text": "",
        }

    def _extract_document(self, ext: str, content: bytes) -> tuple[str, dict[str, Any], list[dict], str, str]:
        """Extract document content, styling, logos, header text, and footer text."""
        suffix = ext or ".tmp"
        tmp_path: Path | None = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(content)
                tmp_path = Path(tmp.name)

            if ext == ".pdf":
                return self._extract_pdf(tmp_path)
            if ext == ".docx":
                return self._extract_docx(tmp_path)
            if ext == ".doc":
                return self._extract_doc(tmp_path)
            return "", self._default_styling(), [], "", ""
        finally:
            if tmp_path:
                tmp_path.unlink(missing_ok=True)

    def _extract_doc(self, path: Path) -> tuple[str, dict[str, Any], list[dict], str, str]:
        """Convert legacy .doc → .docx, then reuse the DOCX extractor (logos/header/footer)."""
        from app.services.doc_converter import DocConversionError, convert_doc_to_docx

        try:
            converted = convert_doc_to_docx(path, output_dir=path.parent)
        except DocConversionError as exc:
            raise FormatExtractionError(str(exc)) from exc
        try:
            return self._extract_docx(converted)
        finally:
            try:
                if converted.exists() and converted != path:
                    converted.unlink(missing_ok=True)
            except Exception:
                pass

    def _extract_pdf(self, path: Path) -> tuple[str, dict[str, Any], list[dict], str, str]:
        try:
            import pdfplumber
        except ImportError as exc:
            raise FormatExtractionError("pdfplumber is required to extract PDF formats") from exc

        pages: list[str] = []
        font_names: list[str] = []
        font_sizes: list[float] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
                
                for char in page.chars[:300]:
                    if char.get("fontname"):
                        font_names.append(str(char["fontname"]))
                    if char.get("size"):
                        font_sizes.append(float(char["size"]))

        # Extract top/bottom page bands for company branding (not candidate bio).
        first_page_lines, last_page_lines = self._pdf_edge_lines(path, pages)
        header_text = self._pick_branding_lines(first_page_lines, region="header")
        footer_text = self._pick_branding_lines(last_page_lines, region="footer")
        if header_text:
            logger.info("Detected PDF header branding: %s", header_text[:120])
        if footer_text:
            logger.info("Detected PDF footer branding: %s", footer_text[:120])

        # Try multiple methods to extract logo (+ optional footer stamp)
        logos = self._extract_logo_from_pdf(path)

        body_size = round(sum(font_sizes) / len(font_sizes), 1) if font_sizes else 11
        font_family = self._clean_font_name(font_names[0]) if font_names else "Helvetica"
        styling = {
            "font_family": font_family,
            "font_size_header": max(14, int(body_size + 3)),
            "font_size_body": int(body_size) or 11,
        }
        return "\n".join(pages), styling, logos, header_text, footer_text
    
    def _extract_logo_from_pdf(self, path: Path) -> list[dict]:
        """Extract logo images from top header area of the first page."""
        logos: list[dict] = []
        logger.info(f"Starting logo extraction from PDF: {path}")
        
        # Method 1: Try PyMuPDF (fitz) for direct image extraction
        try:
            import fitz  # PyMuPDF
            import base64
            
            logger.info("Attempting PyMuPDF extraction...")
            doc = fitz.open(str(path))
            
            if len(doc) > 0:
                page = doc[0]
                page_h = float(page.rect.height or 0)
                header_cutoff = page_h * 0.55 if page_h > 0 else 400.0
                footer_floor = page_h * 0.72 if page_h > 0 else 550.0
                images = page.get_images(full=True)
                logger.info(f"Found {len(images)} images on first page")

                for img_index, img in enumerate(images):
                    try:
                        xref = img[0]
                        rects = page.get_image_rects(xref)
                        if not rects:
                            logger.debug(f"Skipping image {img_index}: no placement rect found")
                            continue
                        top_y = min(float(r.y0) for r in rects)
                        bottom_y = max(float(r.y1) for r in rects)
                        in_header = top_y <= header_cutoff
                        in_footer = bottom_y >= footer_floor
                        if not in_header and not in_footer:
                            logger.debug(
                                "Skipping image %s: outside header/footer (y0=%.1f y1=%.1f)",
                                img_index,
                                top_y,
                                bottom_y,
                            )
                            continue

                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]

                        if len(image_bytes) < 200:
                            logger.debug(f"Skipping small image {img_index}: {len(image_bytes)} bytes")
                            continue
                        if len(image_bytes) > 5_000_000:
                            logger.debug(f"Skipping large image {img_index}: {len(image_bytes)} bytes")
                            continue

                        image_bytes, image_ext = self._normalize_image_bytes(image_bytes, image_ext)
                        if not image_bytes:
                            continue

                        logo_b64 = base64.b64encode(image_bytes).decode("utf-8")
                        position = "footer_center" if in_footer and not in_header else "header_right"
                        source = "pymupdf_footer" if position.startswith("footer") else "pymupdf_header"
                        logos.append({
                            "data": f"data:image/{image_ext};base64,{logo_b64}",
                            "position": position,
                            "width": 150,
                            "height": 75,
                            "source": source,
                            "index": img_index,
                        })
                        logger.info(
                            "Extracted image %s (%s): %s, %s bytes",
                            img_index,
                            source,
                            image_ext,
                            len(image_bytes),
                        )
                        if len(logos) >= 4:
                            break
                    except Exception as img_err:
                        logger.warning(f"Failed to extract image {img_index}: {img_err}")
                        continue

                doc.close()
                if logos:
                    logger.info(f"PyMuPDF extraction successful: {len(logos)} logos")
                    return logos
        except ImportError:
            logger.warning("PyMuPDF (fitz) not available")
        except Exception as e:
            logger.error(f"PyMuPDF extraction failed: {e}")
        
        # Method 2: Try pdfplumber for embedded images
        try:
            import pdfplumber
            import base64
            from io import BytesIO
            from PIL import Image
            
            logger.info("Attempting pdfplumber image extraction...")
            with pdfplumber.open(str(path)) as pdf:
                if len(pdf.pages) > 0:
                    page = pdf.pages[0]
                    try:
                        im = page.to_image(resolution=150)
                        # Try to detect images in page objects
                        if hasattr(page, 'objects') and 'image' in page.objects:
                            for img_obj in page.objects['image'][:3]:
                                try:
                                    # Crop the image region
                                    bbox = (img_obj['x0'], img_obj['top'], 
                                           img_obj['x1'], img_obj['bottom'])
                                    cropped = im.original.crop(bbox)
                                    buffer = BytesIO()
                                    cropped.save(buffer, format='PNG')
                                    logo_b64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                                    logos.append({
                                        "data": f"data:image/png;base64,{logo_b64}",
                                        "position": "detected",
                                        "width": 150,
                                        "height": 75,
                                        "source": "pdfplumber"
                                    })
                                    logger.info(f"pdfplumber extracted image from bbox: {bbox}")
                                except Exception as crop_err:
                                    logger.debug(f"Failed to crop image: {crop_err}")
                    except Exception as page_err:
                        logger.warning(f"pdfplumber page processing failed: {page_err}")
                        
            if logos:
                logger.info(f"pdfplumber extraction successful: {len(logos)} logos")
                return logos
        except ImportError:
            logger.warning("pdfplumber image extraction not available")
        except Exception as e:
            logger.error(f"pdfplumber extraction failed: {e}")
        
        # Method 3: Screenshot-based extraction with better cropping
        try:
            from pdf2image import convert_from_path
            import base64
            from io import BytesIO
            from PIL import Image
            
            logger.info("Attempting pdf2image screenshot extraction...")
            pil_images = convert_from_path(str(path), first_page=1, last_page=1, dpi=200)
            
            if pil_images:
                img = pil_images[0]
                width, height = img.size
                logger.info(f"PDF page size: {width}x{height}")
                
                # Multiple crop strategies for logo detection
                crop_regions = [
                    # Top-right (most common for logos)
                    (width * 0.5, 0, width, height * 0.3, "top_right_wide"),
                    # Narrower top-right
                    (width * 0.7, 0, width, height * 0.2, "top_right_narrow"),
                    # Top-left (alternative)
                    (0, 0, width * 0.4, height * 0.25, "top_left"),
                    # Top center
                    (width * 0.3, 0, width * 0.7, height * 0.2, "top_center"),
                ]
                
                for left, top, right, bottom, region_name in crop_regions:
                    try:
                        left, top, right, bottom = int(left), int(top), int(right), int(bottom)
                        if right > left and bottom > top:
                            cropped = img.crop((left, top, right, bottom))
                            buffer = BytesIO()
                            cropped.save(buffer, format='PNG')
                            logo_b64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                            
                            # Only keep if it has content (not just white space)
                            if len(logo_b64) > 1000:  # Meaningful image
                                logos.append({
                                    "data": f"data:image/png;base64,{logo_b64}",
                                    "position": region_name,
                                    "width": right - left,
                                    "height": bottom - top,
                                    "source": "pdf2image_crop"
                                })
                                logger.info(f"pdf2image extracted from {region_name}: {len(logo_b64)} chars")
                                if len(logos) >= 2:
                                    break
                    except Exception as crop_err:
                        logger.debug(f"Crop {region_name} failed: {crop_err}")
                        continue
                
                if logos:
                    logger.info(f"pdf2image extraction successful: {len(logos)} logos")
                    return logos
        except ImportError:
            logger.warning("pdf2image not available")
        except Exception as e:
            logger.error(f"pdf2image extraction failed: {e}")
        
        logger.warning("No logos extracted from PDF")
        return logos

    def _extract_docx(self, path: Path) -> tuple[str, dict[str, Any], list[dict], str, str]:
        try:
            import docx
        except ImportError as exc:
            raise FormatExtractionError("python-docx and Pillow are required to extract DOCX formats") from exc

        logger.info(f"Extracting DOCX format from: {path}")
        document = docx.Document(str(path))
        lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]

        # Collect logos from every header/footer variant, then body/package fallbacks.
        logos: list[dict] = []
        logos.extend(self._extract_docx_container_images(document, region="header"))
        logos.extend(self._extract_docx_container_images(document, region="footer"))
        if not any(str(logo.get("position") or "").startswith("header") for logo in logos):
            logos.extend(self._extract_docx_top_body_logos(document))
        if not logos:
            logos.extend(self._extract_docx_package_images(document, limit=3))

        # Deduplicate by payload fingerprint while preserving order.
        logos = self._dedupe_logos(logos)

        if logos:
            logger.info("Extracted %s logo(s) from DOCX", len(logos))
        else:
            logger.warning("No logos found in DOCX")

        header_text = self._extract_docx_container_text(document, region="header")
        footer_text = self._extract_docx_container_text(document, region="footer")
        if header_text:
            logger.info("Extracted DOCX header text: %s...", header_text[:120])
        if footer_text:
            logger.info("Extracted DOCX footer text: %s...", footer_text[:120])

        # If no header/footer branding, carefully infer company block from body top lines.
        if not header_text and not footer_text and lines:
            inferred = self._infer_company_header("\n".join(lines[:12]))
            if inferred and inferred.get("lines"):
                header_text = " | ".join(inferred["lines"])
                logger.info("Inferred company branding from DOCX body top lines")

        normal_style = document.styles["Normal"]
        font = normal_style.font
        font_size = int(font.size.pt) if font.size else 11
        styling = {
            "font_family": font.name or "Calibri",
            "font_size_header": max(14, font_size + 3),
            "font_size_body": font_size,
            "font_size_name": max(18, font_size + 8),
            "margin_inches": 0.65,
        }
        logger.info(
            "DOCX extraction complete: %s text lines, %s logos, header=%s, footer=%s",
            len(lines),
            len(logos),
            bool(header_text),
            bool(footer_text),
        )
        return "\n".join(lines), styling, logos, header_text, footer_text

    def _docx_region_containers(self, document: Any, region: str) -> list[Any]:
        """Return header/footer containers including first/even page variants."""
        containers: list[Any] = []
        seen_ids: set[int] = set()
        for section in document.sections:
            candidates: list[Any] = []
            if region == "header":
                candidates = [section.header]
                try:
                    if section.different_first_page_header_footer:
                        candidates.append(section.first_page_header)
                except Exception:
                    pass
                try:
                    candidates.append(section.even_page_header)
                except Exception:
                    pass
            else:
                candidates = [section.footer]
                try:
                    if section.different_first_page_header_footer:
                        candidates.append(section.first_page_footer)
                except Exception:
                    pass
                try:
                    candidates.append(section.even_page_footer)
                except Exception:
                    pass
            for container in candidates:
                if container is None:
                    continue
                cid = id(container)
                if cid in seen_ids:
                    continue
                seen_ids.add(cid)
                containers.append(container)
        return containers

    def _extract_docx_container_text(self, document: Any, region: str) -> str:
        """Extract visible text from header/footer paragraphs and tables."""
        parts: list[str] = []
        for container in self._docx_region_containers(document, region):
            try:
                for para in container.paragraphs:
                    text = re.sub(r"\s+", " ", str(para.text or "")).strip()
                    if text:
                        parts.append(text)
                for table in getattr(container, "tables", []) or []:
                    for row in table.rows:
                        cells = [
                            re.sub(r"\s+", " ", str(cell.text or "")).strip()
                            for cell in row.cells
                            if str(cell.text or "").strip()
                        ]
                        # Unique adjacent cell duplicates from merged cells.
                        uniq: list[str] = []
                        for cell in cells:
                            if not uniq or uniq[-1] != cell:
                                uniq.append(cell)
                        if uniq:
                            parts.append(" | ".join(uniq))
            except Exception as exc:
                logger.debug("DOCX %s text extract failed: %s", region, exc)
        # Keep short branding-like lines preferentially later via _pick_branding_lines.
        cleaned: list[str] = []
        seen: set[str] = set()
        for part in parts:
            key = part.casefold()
            if key in seen:
                continue
            seen.add(key)
            cleaned.append(part)
        if not cleaned:
            return ""
        # Prefer company-like branding lines when available.
        branded = self._pick_branding_lines(cleaned, region=region)
        if branded:
            return branded
        return " | ".join(cleaned[:4])

    def _extract_docx_container_images(self, document: Any, region: str) -> list[dict]:
        """Extract images from header or footer parts (blips, VML, and related image parts)."""
        logos: list[dict] = []
        seen_rids: set[str] = set()
        position = "header_right" if region == "header" else "footer_center"
        source = f"docx_{region}"

        for container in self._docx_region_containers(document, region):
            try:
                element = container._element
            except Exception:
                continue

            # DrawingML blips (modern Word images).
            try:
                blips = element.xpath(".//*[local-name()='blip']")  # nosec - trusted DOCX XML
            except Exception:
                blips = []
            for blip in blips:
                rid = (
                    blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    or blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link")
                    or blip.get("r:embed")
                    or blip.get("r:link")
                )
                if not rid or rid in seen_rids:
                    continue
                seen_rids.add(rid)
                logo = self._logo_from_docx_rid(
                    document,
                    container,
                    rid,
                    position=position,
                    source=source,
                )
                if logo:
                    logos.append(logo)

            # Legacy VML imagedata.
            try:
                imagedata_nodes = element.xpath(".//*[local-name()='imagedata']")  # nosec
            except Exception:
                imagedata_nodes = []
            for node in imagedata_nodes:
                rid = (
                    node.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                    or node.get("r:id")
                    or node.get("id")
                )
                if not rid or rid in seen_rids:
                    continue
                seen_rids.add(rid)
                logo = self._logo_from_docx_rid(
                    document,
                    container,
                    rid,
                    position=position,
                    source=f"{source}_vml",
                )
                if logo:
                    logos.append(logo)

            # Direct image relationships on the header/footer part (covers missed XML refs).
            try:
                part = container.part
                related = getattr(part, "related_parts", {}) or {}
                for rid, rel_part in related.items():
                    if rid in seen_rids:
                        continue
                    content_type = str(getattr(rel_part, "content_type", "") or "").lower()
                    if "image" not in content_type:
                        continue
                    seen_rids.add(rid)
                    logo = self._logo_from_image_bytes(
                        getattr(rel_part, "blob", None),
                        position=position,
                        source=f"{source}_part",
                        preferred_ext=content_type.split("/")[-1] if "/" in content_type else "png",
                    )
                    if logo:
                        logos.append(logo)
            except Exception as exc:
                logger.debug("DOCX %s related image scan failed: %s", region, exc)

            if len(logos) >= 4:
                break
        return logos

    def _logo_from_docx_rid(
        self,
        document: Any,
        container: Any,
        rid: str,
        position: str,
        source: str,
    ) -> dict | None:
        img_data = None
        preferred_ext = "png"
        for part_owner in (container, document):
            try:
                part = part_owner.part
                related = getattr(part, "related_parts", {}) or {}
                if rid not in related:
                    continue
                rel_part = related[rid]
                img_data = getattr(rel_part, "blob", None)
                content_type = str(getattr(rel_part, "content_type", "") or "")
                if "/" in content_type:
                    preferred_ext = content_type.split("/")[-1]
                break
            except Exception:
                continue
        if not img_data:
            # Last-chance package lookup by relationship id on the main document.
            try:
                package_related = getattr(document.part, "related_parts", {}) or {}
                if rid in package_related:
                    rel_part = package_related[rid]
                    img_data = getattr(rel_part, "blob", None)
            except Exception:
                return None
        return self._logo_from_image_bytes(
            img_data,
            position=position,
            source=source,
            preferred_ext=preferred_ext,
        )

    def _logo_from_image_bytes(
        self,
        img_data: bytes | None,
        position: str,
        source: str,
        preferred_ext: str = "png",
    ) -> dict | None:
        import base64
        from io import BytesIO

        if not img_data or len(img_data) < 200 or len(img_data) > 5_000_000:
            return None

        width, height = 150, 75
        ext = (preferred_ext or "png").lower().replace("jpg", "jpeg")
        try:
            from PIL import Image

            img = Image.open(BytesIO(img_data))
            width, height = img.size
            detected = str((img.format or ext or "PNG")).lower().replace("jpg", "jpeg")
            ext = detected or ext
            aspect = width / height if height else 0
            # Keep permissive bounds so wide Aptino wordmarks still pass.
            if width < 24 or height < 16:
                return None
            if aspect and (aspect < 0.15 or aspect > 20):
                return None
        except Exception:
            # EMF/WMF or odd formats — still try normalize.
            pass

        normalized, ext = self._normalize_image_bytes(img_data, ext)
        if not normalized:
            return None
        logo_b64 = base64.b64encode(normalized).decode("utf-8")
        return {
            "data": f"data:image/{ext};base64,{logo_b64}",
            "position": position,
            "width": width,
            "height": height,
            "source": source,
        }

    def _extract_docx_package_images(self, document: Any, limit: int = 3) -> list[dict]:
        """Fallback: any package-embedded image likely used as branding."""
        logos: list[dict] = []
        try:
            related = getattr(document.part, "related_parts", {}) or {}
            for rel in related.values():
                content_type = str(getattr(rel, "content_type", "") or "").lower()
                if "image" not in content_type:
                    continue
                logo = self._logo_from_image_bytes(
                    getattr(rel, "blob", None),
                    position="header_right",
                    source="docx_package",
                    preferred_ext=content_type.split("/")[-1] if "/" in content_type else "png",
                )
                if logo:
                    logos.append(logo)
                if len(logos) >= limit:
                    break
        except Exception as exc:
            logger.warning("DOCX package image extraction failed: %s", exc)
        return logos

    def _dedupe_logos(self, logos: list[dict]) -> list[dict]:
        deduped: list[dict] = []
        seen: set[str] = set()
        for logo in logos:
            data = str(logo.get("data") or "")
            if not data:
                continue
            # Fingerprint on payload tail to avoid giant set memory; enough for uniqueness.
            key = f"{logo.get('position')}|{len(data)}|{data[-64:]}"
            if key in seen:
                continue
            seen.add(key)
            deduped.append(logo)
        return deduped

    def _extract_docx_header_logos(self, document: Any) -> list[dict]:
        """Backward-compatible wrapper."""
        return self._extract_docx_container_images(document, region="header")

    def _extract_docx_footer_logos(self, document: Any) -> list[dict]:
        """Backward-compatible wrapper."""
        return self._extract_docx_container_images(document, region="footer")

    def _extract_docx_top_body_logos(self, document: Any) -> list[dict]:
        """Fallback: extract images in top body content (header-like area)."""
        logos: list[dict] = []
        seen_rids: set[str] = set()

        def add_from_element(element: Any, source: str) -> None:
            try:
                blips = element.xpath(".//*[local-name()='blip']")  # nosec - trusted DOCX XML
            except Exception:
                return
            for blip in blips:
                rid = (
                    blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    or blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link")
                    or blip.get("r:embed")
                )
                if not rid or rid in seen_rids:
                    continue
                seen_rids.add(rid)
                logo = self._logo_from_docx_rid(
                    document,
                    document,
                    rid,
                    position="header_right",
                    source=source,
                )
                if logo:
                    logos.append(logo)

        try:
            for para in list(document.paragraphs)[:20]:
                add_from_element(para._element, "docx_top_body")
                if len(logos) >= 3:
                    return logos
            for table in list(document.tables)[:6]:
                add_from_element(table._element, "docx_top_table")
                if len(logos) >= 3:
                    return logos
        except Exception as exc:
            logger.warning("DOCX top-body logo extraction failed: %s", exc)
        return logos

    def _infer_sections(self, text: str) -> list[str]:
        sections, _labels = self._infer_sections_with_labels(text)
        return sections

    def _infer_sections_with_labels(self, text: str) -> tuple[list[str], dict[str, str]]:
        """Detect section order and the exact heading labels used in the template."""
        found: list[str] = []
        labels: dict[str, str] = {}
        seen = set()
        for raw_line in (text or "").splitlines()[:250]:
            line = re.sub(r"\s+", " ", raw_line).strip()
            if not line or len(line) > 90:
                continue

            normalized = re.sub(r"[^a-z0-9 ]+", "", line.lower()).strip()
            for canonical, aliases in SECTION_ALIASES.items():
                if canonical in seen:
                    continue
                if any(self._heading_matches(normalized, alias) for alias in aliases):
                    found.append(canonical)
                    seen.add(canonical)
                    clean_label = re.sub(r"[:：]+$", "", line).strip()
                    if clean_label:
                        labels[canonical] = clean_label.upper()

        if "header" not in seen:
            found.insert(0, "header")
            seen.add("header")

        if not found:
            return list(DEFAULT_SECTIONS), {}
        return found, labels

    def _heading_matches(self, line: str, alias: str) -> bool:
        """True only for real headings — not body lines that start with a keyword."""
        alias_norm = re.sub(r"[^a-z0-9 ]+", "", alias.lower()).strip()
        if not alias_norm or not line:
            return False
        if line == alias_norm:
            return True
        # Reject long body content like "Experience in AI coding agent tools..."
        if len(line.split()) > 6 or "," in line or len(line) > 48:
            return False
        if line.startswith(f"{alias_norm} "):
            rest = line[len(alias_norm) :].strip()
            return len(rest.split()) <= 3 and len(rest) <= 36
        return False

    def _build_field_mapping(
        self,
        sections: list[str],
        section_labels: dict[str, str] | None = None,
    ) -> dict[str, str]:
        """Map canonical sections to human-readable titles for the generated DOCX."""
        defaults = {
            "summary": "PROFESSIONAL SUMMARY",
            "experience": "PROFESSIONAL EXPERIENCE",
            "education": "EDUCATION",
            "skills": "TECHNICAL SKILLS",
            "projects": "PROJECTS",
            "certifications": "CERTIFICATIONS",
            "achievements": "ACHIEVEMENTS",
            "languages": "LANGUAGES",
        }
        mapping = {
            "name": "Name",
            "email": "Email",
            "phone": "Phone",
            "location": "Location",
        }
        labels = section_labels or {}
        for section in sections:
            if section == "header":
                continue
            label = labels.get(section) or defaults.get(section) or section.replace("_", " ").title()
            if "." in str(label):
                label = defaults.get(section) or section.replace("_", " ").title()
            mapping[section] = label
        return mapping

    def _company_sign_from_text(self, text: str | None) -> dict[str, Any] | None:
        """Build company footer/header sign from extracted header/footer text."""
        raw = (text or "").strip()
        if not raw:
            return None

        parts = [re.sub(r"\s+", " ", p).strip() for p in re.split(r"\s*\|\s*|\n+", raw) if p.strip()]
        if not parts:
            return None

        joined = " ".join(parts).lower()
        if any(marker in joined for marker in ("professional summary", "work experience", "education")):
            return None

        company_markers = (
            "inc",
            "llc",
            "ltd",
            "corp",
            "company",
            "pvt",
            "solutions",
            "technologies",
            "consulting",
            "www.",
            "http",
            "@",
            "email",
            "blvd",
            "suite",
            "street",
            "avenue",
            "road",
        )
        looks_branded = any(m in joined for m in company_markers) or len(parts) >= 2
        if not looks_branded:
            return None

        lines = parts[:3]
        return {
            "name": lines[0],
            "address": lines[1] if len(lines) > 1 else "",
            "contact": lines[2] if len(lines) > 2 else "",
            "lines": lines,
        }

    def _normalize_image_bytes(self, image_bytes: bytes, ext: str | None = None) -> tuple[bytes | None, str]:
        """Convert embedded images to PNG/JPEG suitable for python-docx."""
        from io import BytesIO

        try:
            from PIL import Image
        except Exception:
            cleaned_ext = (ext or "png").lower().replace("jpg", "jpeg")
            if cleaned_ext in {"png", "jpeg", "gif"}:
                return image_bytes, cleaned_ext
            return None, "png"

        cleaned_ext = (ext or "png").lower().replace("jpg", "jpeg").replace("image/", "")
        try:
            img = Image.open(BytesIO(image_bytes))
            # Force load so truncated streams fail here.
            img.load()
            fmt = str((img.format or cleaned_ext or "PNG")).upper()
            if fmt == "JPG":
                fmt = "JPEG"
            if fmt in {"PNG", "JPEG", "GIF", "WEBP", "BMP", "TIFF"}:
                buffer = BytesIO()
                if fmt in {"JPEG", "BMP", "TIFF", "WEBP"} and img.mode in {"RGBA", "P", "LA"}:
                    background = Image.new("RGB", img.size, (255, 255, 255))
                    rgba = img.convert("RGBA")
                    background.paste(rgba, mask=rgba.split()[-1])
                    background.save(buffer, format="PNG")
                    return buffer.getvalue(), "png"
                save_fmt = "PNG" if fmt in {"GIF", "BMP", "TIFF", "WEBP"} else fmt
                out_img = img.convert("RGB") if save_fmt == "JPEG" and img.mode != "RGB" else img
                out_img.save(buffer, format=save_fmt)
                return buffer.getvalue(), save_fmt.lower()

            buffer = BytesIO()
            if img.mode not in {"RGB", "RGBA"}:
                img = img.convert("RGBA")
            img.save(buffer, format="PNG")
            return buffer.getvalue(), "png"
        except Exception:
            if cleaned_ext in {"png", "jpeg", "gif"}:
                return image_bytes, cleaned_ext
            return None, "png"

    def _infer_layout(self, text: str) -> str:
        lines = [line for line in (text or "").splitlines() if line.strip()]
        short_line_ratio = 0
        if lines:
            short_line_ratio = len([line for line in lines if len(line.strip()) < 35]) / len(lines)
        return "two_column" if short_line_ratio > 0.55 and len(lines) > 20 else "single_column"

    def _infer_company_header(self, text: str) -> dict[str, Any] | None:
        """Detect company branding lines at the top of an uploaded template."""
        lines = [re.sub(r"\s+", " ", line).strip() for line in (text or "").splitlines()]
        lines = [line for line in lines if line][:12]
        if not lines:
            return None

        company_markers = (
            "inc.",
            "llc",
            "ltd",
            "corp",
            "company",
            "blvd",
            "suite",
            "avenue",
            "street",
            "www.",
            "http",
            "email:",
            "@",
        )
        section_markers = (
            "summary",
            "experience",
            "education",
            "skills",
            "professional",
            "objective",
            "projects",
        )

        selected: list[str] = []
        for line in lines:
            lower = line.lower()
            if any(marker in lower for marker in section_markers):
                break
            # Skip likely candidate contact lines that are only email/phone without company cues.
            looks_company = any(marker in lower for marker in company_markers)
            looks_address = bool(re.search(r"\b\d{5}(?:-\d{4})?\b", line)) or "," in line
            if looks_company or (looks_address and len(selected) > 0):
                selected.append(line)
            elif not selected and len(line.split()) <= 5 and line[0].isupper():
                # Possible company name line before address/contact.
                selected.append(line)
            if len(selected) >= 3:
                break

        if len(selected) < 2:
            # Allow a single strong company line (Inc/LLC/www/@/address cues).
            if len(selected) == 1:
                lower = selected[0].lower()
                strong = any(marker in lower for marker in company_markers) or bool(
                    re.search(r"\b\d{5}(?:-\d{4})?\b", selected[0])
                )
                if not strong:
                    return None
            else:
                return None

        return {
            "name": selected[0],
            "address": selected[1] if len(selected) > 1 else "",
            "contact": selected[2] if len(selected) > 2 else "",
            "lines": selected[:3],
        }

    def _pdf_edge_lines(self, path: Path, pages: list[str]) -> tuple[list[str], list[str]]:
        """Collect short lines from the top and bottom of the first/last PDF pages."""
        first_page_lines: list[str] = []
        last_page_lines: list[str] = []

        try:
            import fitz

            doc = fitz.open(str(path))
            if len(doc) > 0:
                page = doc[0]
                page_h = float(page.rect.height or 1)
                blocks = page.get_text("blocks") or []
                top_blocks = []
                bottom_blocks = []
                for block in blocks:
                    if len(block) < 5:
                        continue
                    _x0, y0, _x1, y1, text = block[:5]
                    text = str(text or "").strip()
                    if not text:
                        continue
                    if float(y0) <= page_h * 0.22:
                        top_blocks.append((float(y0), text))
                    if float(y1) >= page_h * 0.78:
                        bottom_blocks.append((float(y0), text))
                top_blocks.sort(key=lambda item: item[0])
                bottom_blocks.sort(key=lambda item: item[0])
                for _, text in top_blocks:
                    first_page_lines.extend(
                        [re.sub(r"\s+", " ", line).strip() for line in text.splitlines() if line.strip()]
                    )
                for _, text in bottom_blocks:
                    last_page_lines.extend(
                        [re.sub(r"\s+", " ", line).strip() for line in text.splitlines() if line.strip()]
                    )
            doc.close()
        except Exception as exc:
            logger.debug("PyMuPDF edge-line extraction unavailable: %s", exc)

        if not first_page_lines and pages:
            first_page_lines = [
                re.sub(r"\s+", " ", line).strip()
                for line in pages[0].splitlines()
                if line.strip()
            ][:8]
        if not last_page_lines and pages:
            source = pages[-1]
            last_page_lines = [
                re.sub(r"\s+", " ", line).strip()
                for line in source.splitlines()
                if line.strip()
            ][-8:]
        return first_page_lines, last_page_lines

    def _pick_branding_lines(self, lines: list[str], region: str = "header") -> str:
        """Keep company-like lines; drop candidate names, section titles, page numbers."""
        if not lines:
            return ""

        section_markers = (
            "summary",
            "experience",
            "education",
            "skills",
            "professional",
            "objective",
            "projects",
            "certification",
            "achievement",
        )
        company_markers = (
            "inc",
            "llc",
            "ltd",
            "corp",
            "company",
            "pvt",
            "solutions",
            "technologies",
            "consulting",
            "www.",
            "http",
            "@",
            "email",
            "blvd",
            "suite",
            "street",
            "avenue",
            "road",
            "confidential",
            "copyright",
            "©",
        )

        selected: list[str] = []
        for line in lines:
            clean = re.sub(r"\s+", " ", line).strip()
            if not clean or len(clean) > 120:
                continue
            lower = clean.lower()
            if clean.isdigit() or re.fullmatch(r"page\s*\d+(\s*of\s*\d+)?", lower):
                continue
            if any(marker in lower for marker in section_markers):
                continue
            looks_company = any(marker in lower for marker in company_markers)
            looks_address = bool(re.search(r"\b\d{5}(?:-\d{4})?\b", clean)) or (
                "," in clean and any(ch.isdigit() for ch in clean)
            )
            short_brand = len(clean.split()) <= 6 and clean[0].isupper()
            if region == "footer":
                if looks_company or looks_address or ("www." in lower or "@" in lower):
                    selected.append(clean)
            else:
                if looks_company or looks_address or short_brand:
                    selected.append(clean)
            if len(selected) >= 3:
                break

        if not selected:
            return ""
        if len(selected) == 1:
            joined = selected[0].lower()
            if not any(marker in joined for marker in company_markers) and not re.search(r"\d", selected[0]):
                return ""
        return " | ".join(selected[:3])

    def _build_format_preview(
        self,
        text: str,
        sections: list[str],
        section_labels: dict[str, str] | None = None,
    ) -> str:
        """Sanitized layout hint for the LLM (headings only — no sample person content)."""
        labels = section_labels or {}
        heading_bits = []
        for section in sections:
            if section == "header":
                continue
            label = labels.get(section) or section.replace("_", " ").upper()
            heading_bits.append(str(label).upper())
        layout_hint = "Section order: " + " → ".join(heading_bits[:10]) if heading_bits else ""

        heading_lines: list[str] = []
        for line in (text or "").splitlines():
            clean = re.sub(r"\s+", " ", line).strip()
            if not clean or len(clean) > 48:
                continue
            letters = re.sub(r"[^A-Za-z]", "", clean)
            if len(letters) < 4:
                continue
            upper_ratio = sum(1 for ch in letters if ch.isupper()) / max(len(letters), 1)
            if upper_ratio >= 0.7 or clean.isupper():
                heading_lines.append(clean.upper())
            if len(heading_lines) >= 10:
                break
        parts = []
        if layout_hint:
            parts.append(layout_hint)
        if heading_lines:
            parts.append("Template headings: " + " | ".join(heading_lines))
        return self._preview(" || ".join(parts))

    def _preview(self, text: str) -> str:
        return re.sub(r"\s+", " ", (text or "").strip())[:700]

    def _clean_font_name(self, font_name: str) -> str:
        cleaned = font_name.split("+")[-1].replace("-", " ").strip()
        return cleaned or "Helvetica"

    def _default_styling(self) -> dict[str, Any]:
        return {
            "font_family": "Calibri",
            "font_size_header": 12,
            "font_size_body": 11,
            "font_size_name": 20,
            "margin_inches": 0.7,
        }
