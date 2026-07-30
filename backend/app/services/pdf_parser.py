"""Utility service for extracting raw text from PDF and Word documents."""
from pathlib import Path

import pdfplumber

try:
    import docx
except ImportError:  # pragma: no cover
    docx = None


def _normalize_line(line: str) -> str:
    return " ".join((line or "").strip().split()).casefold()


def _strip_repeated_header_footer(
    page_texts: list[str], *, top_n: int = 3, bottom_n: int = 3, min_repeats: int = 2
) -> list[str]:
    """Remove repeated header/footer lines that appear on most pages.

    Many resumes repeat the same name/contact line at the top and a footer/page
    marker at the bottom of each page; keeping these confuses parsers by
    duplicating content.
    """
    if not page_texts or len(page_texts) < min_repeats:
        return page_texts

    # Collect top/bottom line candidates across pages.
    top_counts: dict[str, int] = {}
    bottom_counts: dict[str, int] = {}
    top_raw: dict[str, str] = {}
    bottom_raw: dict[str, str] = {}

    for text in page_texts:
        lines = [ln for ln in (text or "").splitlines() if ln.strip()]
        if not lines:
            continue
        for ln in lines[:top_n]:
            key = _normalize_line(ln)
            if not key:
                continue
            top_counts[key] = top_counts.get(key, 0) + 1
            top_raw.setdefault(key, ln)
        for ln in lines[-bottom_n:]:
            key = _normalize_line(ln)
            if not key:
                continue
            bottom_counts[key] = bottom_counts.get(key, 0) + 1
            bottom_raw.setdefault(key, ln)

    # Mark as header/footer if repeated on a majority of pages.
    page_count = max(1, len(page_texts))
    threshold = max(min_repeats, int(page_count * 0.6))

    header_keys = {k for k, c in top_counts.items() if c >= threshold}
    footer_keys = {k for k, c in bottom_counts.items() if c >= threshold}

    if not header_keys and not footer_keys:
        return page_texts

    cleaned_pages: list[str] = []
    for text in page_texts:
        raw_lines = (text or "").splitlines()
        kept: list[str] = []
        for ln in raw_lines:
            key = _normalize_line(ln)
            if key and (key in header_keys or key in footer_keys):
                continue
            kept.append(ln)
        cleaned_pages.append("\n".join(kept).strip())

    return cleaned_pages


def extract_text_from_document(path: str) -> str:
    """Extract text from a supported document type."""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        text = _extract_pdf(path).strip()
    elif ext == ".docx":
        text = _extract_docx(path).strip()
    elif ext == ".doc":
        text = _extract_doc(path).strip()
    else:
        return ""
    return repair_collapsed_spaces(text)


def repair_collapsed_spaces(text: str) -> str:
    """Repair PDF/DOCX extraction that jammed words together (missing spaces)."""
    import re

    from app.services.tech_glossary import (
        protect_tech_and_emails,
        restore_tech_names,
        unprotect_placeholders,
    )

    if not text or len(text) < 24:
        return text
    letters = sum(1 for c in text if c.isalpha())
    spaces = text.count(" ")
    if letters < 24:
        return restore_tech_names(text)
    density = spaces / max(letters, 1)
    force_scan = density < 0.10

    protected, restore_map = protect_tech_and_emails(text)
    repaired = protected
    # Avoid splitting protected placeholders; operate on the rest.
    repaired = re.sub(r"([a-z])([A-Z])", r"\1 \2", repaired)
    # Do NOT split letter/digit around emails/tech (already protected). Still skip digit splits
    # inside placeholders by leaving ⟦KEEP…⟧ alone (no [A-Za-z](\d) match on digits-only keep ids).
    repaired = re.sub(r"([A-Za-z])(\d)", r"\1 \2", repaired)
    repaired = re.sub(r"(\d)([A-Za-z])", r"\1 \2", repaired)
    # Undo accidental splits inside placeholders like ⟦ KEEP 0 ⟧
    repaired = re.sub(r"⟦\s*KEEP\s*(\d+)\s*⟧", r"⟦KEEP\1⟧", repaired)

    lexicon = {
        "the", "and", "with", "for", "from", "into", "using", "based", "ready", "native",
        "system", "systems", "service", "services", "agent", "agents", "model", "models",
        "data", "cloud", "backend", "frontend", "platform", "platforms", "interview",
        "candidate", "resume", "parsing", "evaluation", "workflow", "workflows", "tool",
        "tools", "search", "vector", "database", "databases", "context", "protocol",
        "semantic", "multi", "real", "time", "voice", "question", "generation", "insights",
        "production", "grade", "business", "impact", "strong", "focus", "world",
        "architected", "developed", "designed", "implemented", "built", "improved",
        "enabled", "processed", "optimized", "deployed", "integrated", "containerized",
        "reducing", "achieving", "performing", "engineering", "learning", "language",
        "processing", "specialization", "certified", "program", "committee", "member",
        "mentored", "students", "teaching", "assistant", "championship", "table", "tennis",
        "specializing", "building", "applications", "experienced", "developing",
        "scalable", "solutions", "enterprise", "integrations", "reliability",
        "observability", "graduate", "expertise", "institute", "college", "university",
        "bachelor", "master", "technology", "computer", "science", "python",
        "docker", "kubernetes", "machine", "deep", "natural", "large", "prompt",
        "retrieval", "augmented", "protocol", "engineer", "developer", "technologies",
        "celery", "redis", "intelligent", "matching", "analysis", "structured",
        "outputs", "fallback", "strategies", "guardrails", "routing", "bedrock",
        "transcribe", "polly", "enabling", "automated", "recruiter", "velocity",
        "windsurf", "delivering", "across", "multiple", "document", "sources",
        "measured", "citation", "backed", "answer", "financial", "records", "cleaning",
        "feature", "imputation", "outlier", "handling", "robust", "classification",
        "streamlit", "experiment", "tracking", "loan", "conversational", "assistant",
        "external", "capabilities", "interactive", "experiences", "kharagpur",
        "leadership", "achievements", "certifications", "conference",
    }

    def _segment_token(word: str) -> str:
        original = word
        if "⟦KEEP" in original:
            return original
        w = word.lower()
        if len(w) < 10:
            return original
        if any(ch.isdigit() for ch in w) and len(w) < 16:
            return original
        parts: list[str] = []
        i = 0
        while i < len(w):
            match = None
            for j in range(min(len(w), i + 24), i + 3, -1):
                piece = w[i:j]
                if piece in lexicon:
                    match = piece
                    break
            if match:
                parts.append(match)
                i += len(match)
            else:
                if parts and parts[-1] not in lexicon:
                    parts[-1] += w[i]
                else:
                    parts.append(w[i])
                i += 1
        lexicon_hits = sum(1 for p in parts if p in lexicon)
        if lexicon_hits < 2 or len(parts) <= 1:
            return original
        return " ".join(parts)

    out_lines: list[str] = []
    changed = False
    for line in repaired.splitlines():
        if not line.strip():
            out_lines.append(line)
            continue
        # Never aggressively rewrite pure contact lines.
        if "@" in line or "⟦KEEP" in line:
            out_lines.append(line)
            continue
        line_letters = sum(1 for c in line if c.isalpha())
        dens = line.count(" ") / max(line_letters, 1)
        long_tokens = re.findall(r"[A-Za-z]{18,}", line)
        needs_fix = dens < 0.10 or len(long_tokens) >= 1
        if not needs_fix or line_letters < 18:
            out_lines.append(line)
            continue
        chunks = re.findall(r"[A-Za-z]+|[^A-Za-z]+", line)
        rebuilt = "".join(_segment_token(ch) if ch.isalpha() else ch for ch in chunks)
        rebuilt = re.sub(r"\s+", " ", rebuilt).strip()
        if rebuilt != line.strip():
            changed = True
        out_lines.append(rebuilt)

    if not force_scan and not changed and density >= 0.10:
        result = text
    else:
        result = "\n".join(out_lines)
    result = unprotect_placeholders(result, restore_map)
    return restore_tech_names(result)


def extract_text_from_pdf(path: str) -> str:
    """Backward-compatible alias for PDF text extraction."""
    return extract_text_from_document(path)


def _extract_doc(path: str) -> str:
    """Extract text from legacy .doc by converting to .docx first."""
    try:
        from app.services.doc_converter import DocConversionError, convert_doc_to_docx
    except Exception:
        return ""
    try:
        converted = convert_doc_to_docx(Path(path), output_dir=Path(path).parent)
        return _extract_docx(str(converted)).strip()
    except DocConversionError:
        return ""
    except Exception:
        return ""


def _extract_pdf(path: str) -> str:
    text = _extract_pdfplumber(path)
    if not text.strip():
        text = _extract_pymupdf(path)
    return text


def _extract_pdfplumber(path: str) -> str:
    try:
        with pdfplumber.open(path) as pdf:
            pages: list[str] = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            pages = _strip_repeated_header_footer(pages)
            return "\n\n".join(pages)
    except Exception:
        return ""


def _extract_pymupdf(path: str) -> str:
    try:
        import fitz  # pymupdf
        doc = fitz.open(path)
        pages = [page.get_text() for page in doc]
        doc.close()
        pages = _strip_repeated_header_footer(pages)
        return "\n\n".join(pages)
    except Exception:
        return ""


def _extract_docx(path: str) -> str:
    if docx is None:
        return ""
    try:
        document = docx.Document(path)
        # DOCX headers/footers are not included in `document.paragraphs`, but
        # many resumes repeat name/contact at the top of each page inside body
        # content. Apply a light de-duplication of repeated leading/trailing lines.
        parts: list[str] = [p.text for p in document.paragraphs if p.text and str(p.text).strip()]
        # Many client/candidate resumes store experience/skills in tables — include them.
        for table in getattr(document, "tables", []) or []:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = " ".join(
                        p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()
                    ).strip()
                    if cell_text:
                        cells.append(cell_text)
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        pages = [text]
        pages = _strip_repeated_header_footer(pages, min_repeats=999)  # no-op for single "page"
        return pages[0] if pages else text
    except Exception:
        return ""
