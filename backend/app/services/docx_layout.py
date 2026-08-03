"""Shared DOCX layout helpers: right-aligned dates/links on the same line."""

from __future__ import annotations

import re
from typing import Any

# Trailing date ranges / years pulled to the right edge.
_DATE_RIGHT_RE = re.compile(
    r"(?P<left>.*?)\s+"
    r"(?P<right>"
    r"(?:"
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"\.?\s+\d{4}"
    r"(?:\s*(?:[-–—]|to)\s*"
    r"(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"\.?\s+\d{4}|Present|Current|Now|Till\s+Date))?"
    r"|"
    r"\d{4}\s*(?:[-–—]|to)\s*(?:\d{4}|Present|Current|Now)"
    r"|"
    r"(?:Present|Current)\b"
    r")"
    r")\s*$",
    re.IGNORECASE,
)

_URL_RIGHT_RE = re.compile(
    r"(?P<left>.*?)\s+(?P<right>(?:https?://|www\.)\S+)\s*$",
    re.IGNORECASE,
)


def usable_content_width(doc: Any):
    """Return usable body width (Twips/EMU object) for right tab stops."""
    from docx.shared import Inches

    try:
        section = doc.sections[0]
        return section.page_width - section.left_margin - section.right_margin
    except Exception:
        return Inches(6.5)


def split_left_right_meta(text: str) -> tuple[str, str]:
    """
    Split 'Company Name  June 2013 to September 2014' or '... https://...' into
    (left_label, right_meta). Returns (text, '') when nothing to pull right.
    """
    raw = re.sub(r"[ \t]+", " ", str(text or "")).strip()
    if not raw:
        return "", ""

    for pattern in (_DATE_RIGHT_RE, _URL_RIGHT_RE):
        match = pattern.match(raw)
        if not match:
            continue
        left = (match.group("left") or "").strip(" -\u2013\u2014|,;")
        right = (match.group("right") or "").strip()
        if left and right and left.lower() != right.lower():
            return left, right
    return raw, ""


def add_bullet_paragraph(
    doc: Any,
    text: str,
    *,
    font_family: str = "Arial",
    body_size: float = 10,
    color: Any = None,
    space_after: float = 2,
) -> Any:
    """Add a bullet line even when the company template has no 'List Bullet' style."""
    from docx.shared import Pt, Inches

    text = str(text or "").strip()
    if not text:
        return None

    try:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
    except Exception:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        prefix = p.add_run("• ")
        prefix.font.name = font_family
        prefix.font.size = Pt(body_size)
        if color is not None:
            prefix.font.color.rgb = color
        run = p.add_run(text)

    p.paragraph_format.space_after = Pt(space_after)
    run.font.name = font_family
    run.font.size = Pt(body_size)
    run.italic = False
    run.underline = False
    if color is not None:
        run.font.color.rgb = color
    return p


def add_left_right_line(
    doc: Any,
    left: str,
    right: str = "",
    *,
    font_family: str = "Arial",
    body_size: float = 10,
    color_text: Any = None,
    color_muted: Any = None,
    usable_width: Any = None,
    bold_left: bool = True,
    space_before: float = 6,
    space_after: float = 0,
) -> Any:
    """
    Add a paragraph with left text and optional right-aligned date/link at the
    page/content edge. Used for experience, education, projects, and any other
    section line that carries a date or URL.
    """
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_TAB_ALIGNMENT

    left = str(left or "").strip()
    right = str(right or "").strip()
    if not left and not right:
        return None

    # If caller stuffed date/url into left only, peel it off.
    if left and not right:
        left, right = split_left_right_meta(left)

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(space_before)
    line.paragraph_format.space_after = Pt(space_after)

    tab_pos = usable_width if usable_width is not None else usable_content_width(doc)
    if right:
        try:
            # Clear existing stops then set a single right-aligned stop at content edge.
            stops = line.paragraph_format.tab_stops
            for existing in list(getattr(stops, "_tab_stops", []) or []):
                try:
                    stops.remove_tab_stop(existing.position)  # type: ignore[attr-defined]
                except Exception:
                    pass
            line.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
        except Exception:
            try:
                line.paragraph_format.tab_stops.add_tab_stop(
                    tab_pos if tab_pos is not None else Inches(6.5),
                    WD_TAB_ALIGNMENT.RIGHT,
                )
            except Exception:
                pass

    if left:
        run = line.add_run(left)
        run.bold = bool(bold_left)
        run.italic = False
        run.underline = False
        run.font.name = font_family
        run.font.size = Pt(body_size)
        if color_text is not None:
            run.font.color.rgb = color_text

    if right:
        line.add_run("\t")
        date_run = line.add_run(right)
        date_run.bold = False
        date_run.italic = False
        date_run.underline = False
        date_run.font.name = font_family
        date_run.font.size = Pt(body_size)
        if color_muted is not None:
            date_run.font.color.rgb = color_muted
        elif color_text is not None:
            date_run.font.color.rgb = color_text

    return line


def summary_to_bullets(value: Any) -> list[str]:
    """Turn summary text/list into bullet lines (never one mashed paragraph)."""
    import ast
    import json

    if value is None:
        return []

    # Unwrap accidental stringified lists from older pipeline stages.
    if isinstance(value, str):
        text = value.strip()
        if text.startswith("[") and text.endswith("]"):
            for loader in (json.loads, ast.literal_eval):
                try:
                    loaded = loader(text)
                except Exception:
                    loaded = None
                if isinstance(loaded, list):
                    value = loaded
                    break

    if isinstance(value, list):
        raw_parts: list[str] = []
        for item in value:
            if isinstance(item, (list, tuple)):
                raw_parts.extend(str(v or "").strip() for v in item)
            else:
                raw_parts.append(str(item or "").strip())
    else:
        text = str(value or "").strip()
        if not text:
            return []
        # Prefer explicit bullets / newlines; else split on sentence boundaries.
        if re.search(r"[•\n\r]|(?:^|\s)[-–—]\s+", text):
            raw_parts = re.split(r"(?:\r?\n)+|[•]+|(?:^|\s)[-–—]\s+", text)
        else:
            # Keep short one-liners intact; split long prose into sentences.
            if len(text) < 160 and text.count(".") <= 1:
                raw_parts = [text]
            else:
                raw_parts = re.split(r"(?<=[.!?])\s+(?=[A-Z(])", text)

    bullets: list[str] = []
    seen: set[str] = set()
    for part in raw_parts:
        clean = re.sub(r"\s+", " ", str(part or "").strip(" \t•\-–—[]'\""))
        if len(clean) < 12:
            continue
        # Skip list-repr debris.
        if clean.startswith("[") or clean.endswith("]") or clean.count("'") >= 2 and len(clean) < 40:
            continue
        key = clean.lower()
        if key in seen:
            continue
        seen.add(key)
        bullets.append(clean)

    # If over-splitting produced many tiny fragments, keep as fewer richer bullets.
    if len(bullets) >= 8 and sum(len(b) for b in bullets) / max(len(bullets), 1) < 60:
        rejoined = " ".join(bullets)
        bullets = [
            s.strip()
            for s in re.split(r"(?<=[.!?])\s+(?=[A-Z(])", rejoined)
            if len(s.strip()) >= 12
        ] or ([rejoined] if rejoined else [])
    return bullets
