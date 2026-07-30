from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape
import json
import re

# DOCX imports will be loaded dynamically in _render_docx

try:
    from langsmith import traceable
except ImportError:  # pragma: no cover
    def traceable(*args, **kwargs):
        def decorator(func):
            return func
        return decorator

from app.agents.prompts.templates import (
    CLIENT_RESUME_GENERATOR_PROMPT,
    CLIENT_RESUME_GENERATOR_SYSTEM,
    CLIENT_RESUME_GENERATOR_PROMPT_COMPACT,
    CLIENT_RESUME_GENERATOR_SYSTEM_COMPACT,
    CLIENT_RESUME_REWRITER_PROMPT,
    CLIENT_RESUME_REWRITER_SYSTEM,
)
from app.agents.resume_generation_agent import build_resume_document, normalize_resume_document
from app.agents.tools.llm_client import llm_call_json_with_metrics
from app.models.candidate import Candidate
from app.models.client_format import ClientFormat
from app.services.aptino_template import get_aptino_company_footer_lines, is_aptino_template
from app.services.s3_service import download_to_local_path, sanitize_filename, upload_bytes_to_key
from app.services.detailed_resume_parser import parse_resume_detailed
from app.services.pdf_parser import extract_text_from_document
from app.services.structured_resume_store import (
    apply_store_to_candidate_data,
    build_structured_resume,
    document_from_store,
    is_large_resume,
    store_needs_llm_polish,
)
from app.services.resume_page_fitter import estimate_pages, fit_store_to_pages, needs_page_fit
from app.services.resume_llm_condense import condense_store_with_llm, polish_store_with_llm
from app.services.resume_section_quality import (
    audit_and_repair_document,
    critical_findings,
    findings_as_feedback,
)
from app.services.resume_quality_agent import run_section_quality_agents
from app.core.config import settings
from app.core.logging import logger


class ResumeGenerationError(RuntimeError):
    pass


class ResumeGenerationService:
    @traceable(run_type="chain", tags=["resume", "generation", "client-format"])
    def generate(self, candidate: Candidate, client_format: ClientFormat) -> bytes:
        logger.info(
            "starting_resume_generation",
            candidate_id=str(candidate.id),
            candidate_name=candidate.name,
            client_format_id=str(client_format.id),
            client_id=client_format.client_id,
        )
        
        # Use detailed parser for comprehensive resume data
        detailed_data = self._get_detailed_resume_data(candidate)
        # Persist section store on the candidate for reuse / debugging.
        store = build_structured_resume(detailed_data)
        detailed_data = apply_store_to_candidate_data(detailed_data, store)
        existing = dict(candidate.extracted_data or {})
        candidate.extracted_data = {
            **existing,
            "structured_resume": store,
            "raw_text": existing.get("raw_text") or detailed_data.get("raw_resume_text") or "",
            "resume_text": existing.get("resume_text") or detailed_data.get("raw_resume_text") or "",
        }
        logger.info(
            "resume_data_prepared",
            candidate_id=str(candidate.id),
            experience_count=len(detailed_data.get("experience", [])),
            education_count=len(detailed_data.get("education", [])),
            skills_count=len(detailed_data.get("skills", [])),
            experience_bullets=(store.get("stats") or {}).get("experience_bullets"),
            large_resume=is_large_resume(store),
        )
        
        document = self._generate_professional_document(detailed_data, client_format.format_metadata)
        
        # Reuse visual branding from the uploaded client format (logo + company sign).
        # Never copy sample candidate body text from the template.
        format_metadata = dict(client_format.format_metadata or {})
        logos = self._valid_logos(format_metadata.get("logos", []))
        if not logos:
            # Built-in Aptino assets when using Aptino template OR uploaded Aptino-branded format
            # where image extraction failed (common with some DOCX header embeddings).
            branding_blob = " ".join(
                [
                    str(format_metadata.get("template_id") or ""),
                    str(format_metadata.get("template_name") or ""),
                    str((format_metadata.get("company_footer") or {}).get("name") or ""),
                    str((format_metadata.get("company_header") or {}).get("name") or ""),
                    " ".join(str(x) for x in ((format_metadata.get("company_footer") or {}).get("lines") or [])),
                    " ".join(str(x) for x in ((format_metadata.get("company_header") or {}).get("lines") or [])),
                ]
            ).lower()
            if is_aptino_template(format_metadata) or "aptino" in branding_blob:
                from app.services.aptino_template import get_aptino_default_metadata

                logos = self._valid_logos(get_aptino_default_metadata().get("logos", []))
                logger.info(
                    "logo_fallback_aptino_assets",
                    reason="missing_extracted_logos",
                    aptino_template=is_aptino_template(format_metadata),
                )

        # Persist company sign for the DOCX footer (never place in page header).
        company_footer = format_metadata.get("company_footer")
        if not (isinstance(company_footer, dict) and company_footer.get("lines")):
            company_footer = format_metadata.get("company_header")
        if isinstance(company_footer, dict) and company_footer.get("lines"):
            document["company_footer_lines"] = list(company_footer["lines"])
        elif is_aptino_template(format_metadata):
            document["company_footer_lines"] = get_aptino_company_footer_lines()
        else:
            document["company_footer_lines"] = []
        # Backward-compatible alias used by older callers/tests.
        document["company_header_lines"] = list(document["company_footer_lines"])

        header_logos, footer_logos = self._split_logos(logos)
        logo_count = len(header_logos) + len(footer_logos)
        logger.info(
            "client_format_metadata",
            client_format_id=str(getattr(client_format, "id", "")),
            logo_count=logo_count,
            header_logo_count=len(header_logos),
            footer_logo_count=len(footer_logos),
            has_logos=logo_count > 0,
            template_id=format_metadata.get("template_id"),
            has_company_footer=bool(document.get("company_footer_lines")),
        )

        if logo_count > 0:
            for i, logo in enumerate(header_logos + footer_logos):
                data_size = len(logo.get("data", "")) if logo.get("data") else 0
                logger.info(
                    f"logo_{i+1}_info",
                    position=logo.get("position"),
                    source=logo.get("source"),
                    data_size=data_size,
                )

        document["client_header_text"] = ""
        document["client_footer_text"] = ""

        # Generate DOCX with the client logo + company sign (editable format)
        try:
            docx_bytes = self._render_docx(document, format_metadata, header_logos, footer_logos)
            logger.info(
                "docx_rendered_successfully",
                candidate_id=str(candidate.id),
                docx_size=len(docx_bytes),
                logo_included=logo_count > 0,
            )
        except Exception as e:
            logger.error(
                "docx_rendering_failed",
                candidate_id=str(candidate.id),
                error=str(e),
                exc_info=True,
            )
            raise ResumeGenerationError(f"Failed to render DOCX: {e}") from e

        safe_client = sanitize_filename(getattr(client_format, "client_id", None) or "client")
        object_key = f"resumes/{candidate.recruiter_id}/generated/{candidate.id}_{safe_client}.docx"

        logger.info(
            "uploading_generated_resume",
            candidate_id=str(candidate.id),
            object_key=object_key,
        )

        upload_bytes_to_key(
            recruiter_id=str(candidate.recruiter_id),
            object_key=object_key,
            content=docx_bytes,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        return docx_bytes

    @traceable(run_type="retriever", tags=["resume", "data-gathering"])
    def _get_detailed_resume_data(self, candidate: Candidate) -> dict[str, Any]:
        """
        Get comprehensive resume data using multiple sources.
        Ensures ALL candidate content is used for generation.
        """
        logger.info(
            "gathering_resume_data",
            candidate_id=str(candidate.id),
            has_extracted_data=bool(candidate.extracted_data),
            has_resume_path=bool(candidate.resume_path),
        )

        # Try to get resume text - from stored text first
        resume_text = None
        if candidate.extracted_data:
            resume_text = (
                candidate.extracted_data.get("raw_text")
                or candidate.extracted_data.get("resume_text")
                or candidate.extracted_data.get("full_text")
            )

        # If no stored text, extract from the original resume file.
        if not resume_text and candidate.resume_path:
            try:
                resume_text = self._extract_resume_text_from_path(candidate)
                logger.info(
                    "extracted_resume_text_from_file",
                    candidate_id=str(candidate.id),
                    text_length=len(resume_text),
                )
            except Exception as e:
                logger.error("failed_to_extract_resume_text", candidate_id=str(candidate.id), error=str(e))
        
        # Parse with detailed parser if we have text
        detailed_data = {}
        if resume_text:
            try:
                detailed_data = parse_resume_detailed(resume_path=None, resume_text=resume_text)
                logger.info(
                    "detailed_parse_complete",
                    candidate_id=str(candidate.id),
                    experience_count=len(detailed_data.get("experience", [])),
                    education_count=len(detailed_data.get("education", [])),
                    skills_count=len(detailed_data.get("skills", [])),
                )
            except Exception as e:
                logger.error("detailed_parse_failed", candidate_id=str(candidate.id), error=str(e))
        else:
            logger.warning("no_resume_text_available", candidate_id=str(candidate.id))
        
        # Get pipeline extracted data as fallback
        pipeline = candidate.extracted_data or {}
        manual_entry = pipeline.get("manual_entry") if isinstance(pipeline.get("manual_entry"), dict) else {}
        
        # MERGE STRATEGY: Use detailed data first, fallback to pipeline data, never lose content
        def merge_field(detailed_key: str, pipeline_keys: list[str], default=None):
            """Get field from detailed data or pipeline with multiple fallback keys."""
            # Try detailed data first
            value = detailed_data.get(detailed_key)
            if value:
                return value
            
            # Try all pipeline keys
            for key in pipeline_keys:
                value = pipeline.get(key)
                if value:
                    return value
            
            return default or []
        
        # Experience: merge ALL sources so partial parsers cannot drop older roles.
        experience = self._merge_experience_lists(
            detailed_data.get("experience", []),
            pipeline.get("experience"),
            pipeline.get("resume_experience"),
            pipeline.get("work_experience"),
            pipeline.get("work_history"),
            self._employment_history_from_candidate(candidate),
        )
        if not experience:
            fallback_experience = self._fallback_experience_from_candidate_fields(candidate)
            if fallback_experience:
                experience = [fallback_experience]
        
        # Education: merge all sources so partial parsers cannot drop degrees.
        education = self._merge_named_record_lists(
            detailed_data.get("education", []),
            pipeline.get("education"),
            pipeline.get("resume_education"),
            pipeline.get("degrees"),
            pipeline.get("academic_background"),
            self._education_history_from_candidate(candidate),
            key_fields=("institution", "school", "university", "degree", "title"),
        )
        if not education and resume_text:
            education = self._extract_education_from_text(resume_text)
        
        # Skills: comprehensive merge from all sources
        skills = detailed_data.get("skills", [])
        if not skills:
            # Combine all skill sources
            all_skills: list[str] = []
            for source in [
                candidate.skills_matched,
                candidate.skills_not_matched,
                self._candidate_attr(candidate, "primary_skills"),
                self._candidate_attr(candidate, "secondary_skills"),
                self._candidate_attr(candidate, "other_skills"),
                pipeline.get("skills_matched"),
                pipeline.get("skills_not_matched"),
                pipeline.get("skills"),
                pipeline.get("technical_skills"),
                pipeline.get("core_skills"),
            ]:
                all_skills.extend(self._extract_skill_items(source))
            if not all_skills and resume_text:
                all_skills.extend(self._extract_skills_from_text(resume_text))
            skills = all_skills

        skills = self._clean_skill_list(skills)
        # Prefer parser category labels (Tools / Platforms / ...) over heuristic regrouping.
        parser_categories = detailed_data.get("skills_by_category")
        if isinstance(parser_categories, dict) and parser_categories:
            skills_by_category = {
                str(k).strip(): self._clean_skill_list(v)
                for k, v in parser_categories.items()
                if str(k).strip() and self._clean_skill_list(v)
            }
        else:
            skills_by_category = self._group_skills_for_resume(skills)
        summary = (
            detailed_data.get("summary")
            or detailed_data.get("professional_summary")
            or pipeline.get("summary")
            or pipeline.get("professional_summary")
            or pipeline.get("objective")
            or self._candidate_attr(candidate, "main_summary")
            or self._candidate_attr(candidate, "linkedin_summary")
            or manual_entry.get("additional_info")
            or manual_entry.get("notes")
            or self._fallback_summary_from_candidate_fields(candidate, skills)
            or ""
        )
        languages = merge_field("languages", ["languages", "spoken_languages"])
        candidate_language = self._candidate_attr(candidate, "language")
        if candidate_language:
            languages = self._dedupe_text([
                *[self._clean_inline_text(item) for item in self._as_list(languages)],
                *[self._clean_inline_text(item) for item in self._as_list(candidate_language)],
            ])
        
        # Build comprehensive result with ALL fields
        result = {
            # Contact info
            "name": (
                detailed_data.get("name")
                or self._candidate_attr(candidate, "name")
                or self._candidate_attr(candidate, "display_name")
                or pipeline.get("name")
                or ""
            ),
            "email": detailed_data.get("email") or self._candidate_attr(candidate, "email") or pipeline.get("email") or "",
            "phone": detailed_data.get("phone") or self._candidate_attr(candidate, "phone") or pipeline.get("phone") or "",
            "location": (
                detailed_data.get("location")
                or self._candidate_attr(candidate, "location")
                or pipeline.get("location")
                or self._candidate_location(candidate)
                or ""
            ),
            "linkedin": (
                detailed_data.get("linkedin") 
                or pipeline.get("linkedin") 
                or pipeline.get("linkedin_url")
                or self._candidate_attr(candidate, "linkedin_url")
                or ""
            ),
            "portfolio": (
                detailed_data.get("portfolio")
                or pipeline.get("portfolio")
                or pipeline.get("website")
                or self._candidate_attr(candidate, "web_path")
                or ""
            ),
            "github": detailed_data.get("github") or pipeline.get("github") or "",
            
            # Job context
            "job_applied": self._candidate_attr(candidate, "job_applied"),
            "job_role": (
                self._candidate_attr(candidate, "job_role")
                or self._candidate_attr(candidate, "job_title")
                or pipeline.get("target_role")
                or ""
            ),
            "client_name": self._candidate_attr(candidate, "client_name") or pipeline.get("client_name") or "",
            
            # Content sections - comprehensive merge
            "summary": summary,
            "skills": skills,
            "skills_by_category": skills_by_category,
            "experience": experience,
            "education": education,
            "projects": merge_field("projects", ["projects", "personal_projects", "open_source"]),
            "certifications": merge_field("certifications", ["certifications", "certificates", "licenses"]),
            "achievements": merge_field("achievements", ["achievements", "awards", "honors", "accomplishments"]),
            "languages": languages,
            "publications": merge_field("publications", ["publications", "papers", "research"]),
            "interests": merge_field("interests", ["interests", "hobbies"]),
            # Canonical section-wise JSON passed to LLM for consistent mapping.
            "structured_resume_json": {
                "summary": (
                    summary
                ),
                "skills": skills,
                "skills_by_category": skills_by_category,
                "experience": experience,
                "education": education,
                "projects": merge_field("projects", ["projects", "personal_projects", "open_source"]),
                "certifications": merge_field("certifications", ["certifications", "certificates", "licenses"]),
                "achievements": merge_field("achievements", ["achievements", "awards", "honors", "accomplishments"]),
                "languages": languages,
            },
            
            # Raw data for reference
            "extracted_data": detailed_data,
            "pipeline_data": pipeline,
            "raw_resume_text": resume_text or "",
        }
        
        # Log summary of what we found
        logger.info(
            "resume_data_merged",
            candidate_id=str(candidate.id),
            has_summary=bool(result["summary"]),
            skills_count=len(result["skills"]),
            experience_count=len(result["experience"]),
            education_count=len(result["education"]),
            projects_count=len(result["projects"]),
            certifications_count=len(result["certifications"]),
        )
        
        return result

    def _candidate_attr(self, candidate: Candidate, name: str, default: Any = "") -> Any:
        """Safely read SQLAlchemy/model attributes without triggering generation failure."""
        try:
            value = getattr(candidate, name, default)
        except Exception:
            return default
        return default if value is None else value

    def _candidate_location(self, candidate: Candidate) -> str:
        parts = [
            self._candidate_attr(candidate, "city"),
            self._candidate_attr(candidate, "state"),
            self._candidate_attr(candidate, "country"),
        ]
        return ", ".join(str(part).strip() for part in parts if str(part or "").strip())

    def _history_items(self, candidate: Candidate, attr: str) -> list[Any]:
        try:
            value = getattr(candidate, attr, []) or []
        except Exception:
            return []
        return list(value) if isinstance(value, (list, tuple)) else []

    def _format_history_date(self, value: Any) -> str:
        if not value:
            return ""
        if hasattr(value, "strftime"):
            return value.strftime("%b %Y")
        return self._clean_inline_text(value)

    def _date_range(self, start: Any, end: Any) -> str:
        start_text = self._format_history_date(start)
        end_text = self._format_history_date(end)
        if start_text and end_text:
            return f"{start_text} - {end_text}"
        return start_text or end_text

    def _experience_identity(self, item: Any) -> str:
        if not isinstance(item, dict):
            return re.sub(r"\s+", " ", str(item or "").strip().lower())
        company = str(
            item.get("company")
            or item.get("organization")
            or item.get("employer")
            or ""
        ).strip().lower()
        title = str(
            item.get("title")
            or item.get("role")
            or item.get("position")
            or item.get("job_title")
            or ""
        ).strip().lower()
        duration = str(
            item.get("duration")
            or item.get("date")
            or item.get("dates")
            or item.get("year")
            or ""
        ).strip().lower()
        return re.sub(r"\s+", " ", f"{company}|{title}|{duration}")

    def _merge_experience_lists(self, *sources: Any) -> list[dict[str, Any]]:
        """Union experience entries across parsers without dropping older roles."""
        merged: list[dict[str, Any]] = []
        seen: set[str] = set()
        for source in sources:
            for item in self._as_list(source):
                if isinstance(item, dict):
                    clone = dict(item)
                    desc = (
                        clone.get("description")
                        or clone.get("details")
                        or clone.get("responsibilities")
                        or clone.get("achievements")
                        or []
                    )
                    clone["description"] = [str(d).strip() for d in self._as_list(desc) if str(d).strip()]
                    key = self._experience_identity(clone)
                    if not key or key == "||":
                        continue
                    if key in seen:
                        # Prefer the entry with richer bullets when duplicates collide.
                        for idx, existing in enumerate(merged):
                            if self._experience_identity(existing) == key:
                                existing_desc = self._as_list(existing.get("description"))
                                new_desc = self._as_list(clone.get("description"))
                                if len(new_desc) > len(existing_desc):
                                    merged[idx] = clone
                                break
                        continue
                    seen.add(key)
                    merged.append(clone)
                else:
                    text = self._clean_inline_text(item)
                    if not text:
                        continue
                    key = text.lower()
                    if key in seen:
                        continue
                    seen.add(key)
                    merged.append({"title": text, "company": "", "duration": "", "description": []})
        return merged

    def _merge_named_record_lists(
        self,
        *sources: Any,
        key_fields: tuple[str, ...] = ("title", "name"),
    ) -> list[Any]:
        """Union education/project-like records across sources with light de-dupe."""
        merged: list[Any] = []
        seen: set[str] = set()
        for source in sources:
            for item in self._as_list(source):
                if isinstance(item, dict):
                    parts = [str(item.get(field) or "").strip().lower() for field in key_fields]
                    key = re.sub(r"\s+", " ", "|".join(parts)).strip("|")
                    if not key:
                        key = re.sub(r"\s+", " ", str(item).lower())
                    if key in seen:
                        continue
                    seen.add(key)
                    merged.append(dict(item))
                else:
                    text = self._clean_inline_text(item)
                    if not text:
                        continue
                    key = text.lower()
                    if key in seen:
                        continue
                    seen.add(key)
                    merged.append(text)
        return merged

    def _employment_history_from_candidate(self, candidate: Candidate) -> list[dict[str, Any]]:
        entries: list[dict[str, Any]] = []
        for row in self._history_items(candidate, "employment_history"):
            title = self._candidate_attr(row, "job_title")
            company = self._candidate_attr(row, "company_name")
            duration = self._date_range(
                self._candidate_attr(row, "start_date"),
                self._candidate_attr(row, "end_date"),
            )
            salary = self._candidate_attr(row, "salary")
            details = [f"Salary: {salary}"] if salary else []
            if title or company or duration or details:
                entries.append(
                    {
                        "title": title or "Professional Experience",
                        "company": company,
                        "duration": duration,
                        "description": details,
                    }
                )
        return entries

    def _education_history_from_candidate(self, candidate: Candidate) -> list[dict[str, Any]]:
        entries: list[dict[str, Any]] = []
        for row in self._history_items(candidate, "education_history"):
            degree_parts = [
                self._candidate_attr(row, "degree_type"),
                self._candidate_attr(row, "degree_name"),
            ]
            degree = " ".join(str(part).strip() for part in degree_parts if str(part or "").strip())
            major = self._candidate_attr(row, "major")
            institution = self._candidate_attr(row, "school_name")
            year = (
                self._format_history_date(self._candidate_attr(row, "degree_date"))
                or self._date_range(
                    self._candidate_attr(row, "start_date"),
                    self._candidate_attr(row, "end_date"),
                )
            )
            details = []
            if major:
                details.append(f"Major: {major}")
            school_type = self._candidate_attr(row, "school_type")
            if school_type:
                details.append(f"School type: {school_type}")
            if degree or institution or year or details:
                entries.append(
                    {
                        "degree": degree or "Education",
                        "institution": institution,
                        "year": year,
                        "details": details,
                    }
                )
        return entries

    def _fallback_experience_from_candidate_fields(self, candidate: Candidate) -> dict[str, Any] | None:
        title = (
            self._candidate_attr(candidate, "job_title")
            or self._candidate_attr(candidate, "job_role")
            or self._candidate_attr(candidate, "job_applied")
        )
        company = self._candidate_attr(candidate, "company_name")
        total_exp = self._candidate_attr(candidate, "total_exp")
        us_exp = self._candidate_attr(candidate, "us_exp")
        if not any([title, company, total_exp, us_exp]):
            return None

        details = []
        if total_exp:
            details.append(f"Total professional experience: {total_exp}.")
        if us_exp:
            details.append(f"US experience: {us_exp}.")
        return {
            "title": title or "Professional Experience",
            "company": company,
            "description": details,
        }

    def _fallback_summary_from_candidate_fields(self, candidate: Candidate, skills: list[str]) -> str:
        role = (
            self._candidate_attr(candidate, "job_title")
            or self._candidate_attr(candidate, "job_role")
            or self._candidate_attr(candidate, "job_applied")
            or "Professional"
        )
        total_exp = self._candidate_attr(candidate, "total_exp")
        parts = []
        if role:
            if total_exp:
                parts.append(f"{role} with {total_exp} years of professional experience")
            else:
                parts.append(f"{role} with a background aligned to the target role")
        if skills:
            parts.append(f"skilled in {', '.join(skills[:8])}")
        if not parts:
            return ""
        summary = "; ".join(parts).strip()
        return summary[0].upper() + summary[1:] + "."

    def _compress_candidate_data(self, data: dict[str, Any]) -> dict[str, Any]:
        """Trim debug fields but keep enough detail for a professional rewrite.

        Keep nearly all experience/project content so the LLM cannot silently drop roles.
        """
        drop_keys = {"pipeline_data", "extracted_data"}
        compressed: dict[str, Any] = {}
        for key, value in data.items():
            if key in drop_keys:
                continue
            if value is None or value == "" or value == []:
                continue
            compressed[key] = value

        # Keep all roles; only lightly trim extremely long bullets for token budget.
        experience = compressed.get("experience", [])
        if isinstance(experience, list):
            for exp in experience:
                if not isinstance(exp, dict):
                    continue
                desc = exp.get("description", [])
                if isinstance(desc, list):
                    exp["description"] = [str(d)[:600] for d in desc if str(d).strip()]
                if isinstance(exp.get("technologies"), list) and len(exp.get("technologies", [])) > 30:
                    exp["technologies"] = exp["technologies"][:30]

        projects = compressed.get("projects", [])
        if isinstance(projects, list):
            for proj in projects:
                if not isinstance(proj, dict):
                    continue
                desc = proj.get("description") or proj.get("details") or []
                if isinstance(desc, list):
                    proj["description"] = [str(d)[:600] for d in desc if str(d).strip()]

        skills = self._clean_skill_list(compressed.get("skills", []))
        if skills:
            compressed["skills"] = skills[:150]
            grouped_skills = self._group_skills_for_resume(skills)
            if grouped_skills:
                compressed["skills_by_category"] = grouped_skills

        summary = compressed.get("summary", "")
        if len(summary) > 2500:
            compressed["summary"] = summary[:2500]

        # Keep ample raw text so long multi-role resumes can still be recovered.
        raw_text = str(compressed.get("raw_resume_text") or "")
        if raw_text:
            compressed["raw_resume_text"] = raw_text[:40000]

        return compressed

    @traceable(run_type="chain", tags=["resume", "document-generation"])
    def _generate_professional_document(
        self,
        candidate_data: dict[str, Any],
        format_metadata: dict[str, Any] | None,
    ) -> dict[str, Any]:
        """Generate a client-formatted resume from the section store (fast, genuine)."""
        metadata = self._metadata_for_generation(format_metadata or {})

        # 1) Full section store — source of truth, no mixing across sections.
        full_store = build_structured_resume(candidate_data)
        pages_full = estimate_pages(full_store)
        target_pages = float(getattr(settings, "RESUME_TARGET_PAGES", 3.0) or 3.0)

        # 2) Fast path: deterministic bullets/format first.
        #    LLM only when oversized (condense) OR small resume still has mashed/broken text.
        if needs_page_fit(full_store, target_pages=target_pages):
            render_store = fit_store_to_pages(full_store, target_pages=target_pages)
            mode = "store_page_fit"
            if bool(getattr(settings, "RESUME_LLM_CONDENSE", True)):
                condensed = condense_store_with_llm(render_store, target_pages=target_pages)
                if condensed:
                    render_store = condensed
                    mode = "store_page_fit_llm"
        else:
            render_store = full_store
            mode = "store_keep_all"
            # Skip slow polish for clean text and for large resumes (was timing out ~6 min).
            needs_polish = store_needs_llm_polish(render_store)
            if (
                bool(getattr(settings, "RESUME_LLM_CONDENSE", True))
                and needs_polish
                and not is_large_resume(full_store)
            ):
                polished = polish_store_with_llm(render_store)
                if polished:
                    render_store = polished
                    mode = "store_keep_all_llm"
            elif needs_polish and is_large_resume(full_store):
                logger.info(
                    "resume_llm_polish_skipped",
                    reason="large_resume_use_deterministic_bullets",
                    pages=round(pages_full, 2),
                )

        candidate_data = apply_store_to_candidate_data(candidate_data, render_store)
        if isinstance(candidate_data.get("extracted_data"), dict):
            candidate_data["extracted_data"] = {
                **candidate_data["extracted_data"],
                "structured_resume": full_store,
                "structured_resume_render": render_store,
            }

        draft_document = document_from_store(render_store, metadata)
        draft_document = normalize_resume_document(draft_document, candidate_data)
        draft_document["client_name"] = candidate_data.get("client_name", "")
        draft_document = self._enforce_document_reliability(draft_document, candidate_data, metadata)

        # Quality-first: deterministic audit, then section critic/repair agents (max 2 rounds).
        draft_document, section_findings = audit_and_repair_document(draft_document, candidate_data)
        draft_document = self._enforce_document_reliability(draft_document, candidate_data, metadata)
        agent_metrics: dict[str, Any] = {"llm_calls": 0, "sections_repaired": [], "rounds": 0}
        if critical_findings(section_findings) and bool(getattr(settings, "RESUME_LLM_CONDENSE", True)):
            draft_document, section_findings, agent_metrics = run_section_quality_agents(
                draft_document,
                candidate_data,
                max_rounds=2,
                max_llm_calls=5,
            )
            draft_document = self._enforce_document_reliability(draft_document, candidate_data, metadata)
        draft_issues = self._resume_quality_feedback(candidate_data, draft_document)
        draft_issues = list(dict.fromkeys([*findings_as_feedback(section_findings), *draft_issues]))

        logger.info(
            "resume_generation_store_ready",
            mode=mode,
            pages_full=round(pages_full, 2),
            pages_render=round(estimate_pages(render_store), 2),
            target_pages=target_pages,
            experience_count=(render_store.get("stats") or {}).get("experience_count"),
            experience_bullets=(render_store.get("stats") or {}).get("experience_bullets"),
            skills_count=(render_store.get("stats") or {}).get("skills_count"),
            education_count=len(render_store.get("education") or []),
            fit=render_store.get("fit"),
            section_critical=len(critical_findings(section_findings)),
            section_findings=len(section_findings),
            quality_agent_calls=agent_metrics.get("llm_calls"),
            quality_agent_repaired=agent_metrics.get("sections_repaired"),
        )

        best_document = draft_document
        best_issues = draft_issues
        best_findings = section_findings
        llm_used = mode.endswith("_llm") or bool(agent_metrics.get("llm_calls"))

        # Fallback full rewrite only if critical section issues remain after agents.
        needs_quality_repair = bool(critical_findings(section_findings)) and not agent_metrics.get(
            "sections_repaired"
        )
        if needs_quality_repair and bool(getattr(settings, "RESUME_LLM_CONDENSE", True)):
            logger.info(
                "resume_section_quality_repair_started",
                critical=len(critical_findings(section_findings)),
                issues=findings_as_feedback(critical_findings(section_findings))[:8],
            )
            repaired = self._rewrite_document(
                draft_document=best_document,
                candidate_data=candidate_data,
                metadata=metadata,
                feedback=[
                    "CRITICAL SECTION QUALITY REPAIR REQUIRED:",
                    "- Fix ONLY the listed section issues.",
                    "- Do NOT invent employers, schools, skills, dates, or metrics.",
                    "- Experience: real company/title/dates; discrete bullets; never Environment as title.",
                    "- Skills: keep source category names; atomic skill names only (FastAPI not Fast API).",
                    "- Education: Degree + full Institution + Year (no Institute/College stubs).",
                    "- Certifications must be separate items; never mix Leadership/Achievements into Certifications.",
                    "- Projects must be separate entries with their own bullets.",
                    "- Section titles must be canonical (PROFESSIONAL SUMMARY, TECHNICAL SKILLS, etc).",
                    "- Header must include full name and full email.",
                    *findings_as_feedback(critical_findings(section_findings) or section_findings)[:10],
                    *draft_issues[:6],
                ],
            )
            if repaired:
                repaired = self._enforce_document_reliability(repaired, candidate_data, metadata)
                repaired, repaired_findings = audit_and_repair_document(repaired, candidate_data)
                repaired = self._enforce_document_reliability(repaired, candidate_data, metadata)
                repaired_issues = self._resume_quality_feedback(candidate_data, repaired)
                repaired_issues = list(
                    dict.fromkeys([*findings_as_feedback(repaired_findings), *repaired_issues])
                )
                llm_used = True
                if self._document_is_better(
                    repaired, repaired_issues, best_document, best_issues, candidate_data
                ) or (
                    len(critical_findings(repaired_findings)) < len(critical_findings(best_findings))
                ):
                    # Prefer repaired unless it dropped most experience content.
                    if not self._document_is_substantially_thinner(repaired, best_document, candidate_data):
                        best_document = repaired
                        best_issues = repaired_issues
                        best_findings = repaired_findings

            # If still critical on a non-huge resume, one compose pass as last resort.
            if critical_findings(best_findings) and not is_large_resume(full_store):
                composed = self._compose_document_with_llm(
                    candidate_data=candidate_data,
                    metadata=metadata,
                    baseline_document=best_document,
                )
                if composed:
                    composed = self._enforce_document_reliability(composed, candidate_data, metadata)
                    composed, composed_findings = audit_and_repair_document(composed, candidate_data)
                    composed = self._enforce_document_reliability(composed, candidate_data, metadata)
                    composed_issues = self._resume_quality_feedback(candidate_data, composed)
                    composed_issues = list(
                        dict.fromkeys([*findings_as_feedback(composed_findings), *composed_issues])
                    )
                    llm_used = True
                    if not self._document_is_substantially_thinner(composed, best_document, candidate_data):
                        if len(critical_findings(composed_findings)) <= len(critical_findings(best_findings)):
                            best_document = composed
                            best_issues = composed_issues
                            best_findings = composed_findings

        # Final mandatory section audit before DOCX.
        best_document, final_findings = audit_and_repair_document(best_document, candidate_data)
        best_document = self._enforce_document_reliability(best_document, candidate_data, metadata)
        final_critical = critical_findings(final_findings)
        if final_critical:
            logger.warning(
                "resume_section_quality_remaining_critical",
                count=len(final_critical),
                issues=findings_as_feedback(final_critical)[:8],
            )
        else:
            logger.info("resume_section_quality_passed", warnings=len(final_findings))

        logger.info(
            "resume_quality_document_selected",
            mode=mode,
            baseline_issues=len(draft_issues),
            final_issues=len(best_issues),
            section_critical_remaining=len(final_critical),
            critical_remaining=bool(final_critical) or self._has_critical_quality_issues(best_issues),
            llm_compose_used=llm_used,
        )
        return best_document

    def _polish_summary_only(self, candidate_data: dict[str, Any], summary: str) -> str:
        """Optional light polish for summary text only — never touches experience/skills."""
        text = str(summary or "").strip()
        if len(text) < 80:
            return text
        try:
            from app.agents.tools.llm_client import llm_call

            polished = llm_call(
                "You polish professional resume summaries. Preserve all facts, technologies, "
                "domains, and years of experience. Do not invent content. Return plain text only.",
                (
                    "Polish this professional summary for clarity and impact while keeping EVERY fact:\n\n"
                    f"{text[:5000]}\n\n"
                    f"Target role: {candidate_data.get('job_role') or candidate_data.get('job_title') or ''}"
                ),
                max_tokens=1200,
            )
            cleaned = re.sub(r"\s+", " ", str(polished or "").strip())
            if cleaned and len(cleaned) >= max(120, int(len(text) * 0.55)):
                return cleaned
        except Exception as exc:
            logger.warning("summary_polish_failed", error=str(exc))
        return text

    def _experience_bullet_count(self, document: dict[str, Any]) -> int:
        total = 0
        for section in document.get("sections") or []:
            if not isinstance(section, dict):
                continue
            if self._canonical_resume_section(section.get("title") or section.get("type")) != "experience":
                continue
            for item in self._as_list(section.get("content")):
                if not isinstance(item, dict):
                    continue
                desc = item.get("description") or item.get("bullets") or item.get("responsibilities") or []
                total += len(self._as_list(desc))
        return total

    def _document_is_substantially_thinner(
        self,
        candidate_doc: dict[str, Any],
        baseline_doc: dict[str, Any],
        candidate_data: dict[str, Any],
    ) -> bool:
        """True when LLM output dropped most roles/bullets vs grounded baseline."""
        src_roles = len(self._as_list(candidate_data.get("experience")))
        base_roles = self._experience_role_count(baseline_doc)
        cand_roles = self._experience_role_count(candidate_doc)
        expected_roles = max(src_roles, base_roles)
        if expected_roles >= 3 and cand_roles < max(2, int(expected_roles * 0.6)):
            return True

        base_bullets = self._experience_bullet_count(baseline_doc)
        cand_bullets = self._experience_bullet_count(candidate_doc)
        if base_bullets >= 12 and cand_bullets < max(6, int(base_bullets * 0.45)):
            return True
        return False

    def _experience_role_count(self, document: dict[str, Any]) -> int:
        for section in document.get("sections") or []:
            if not isinstance(section, dict):
                continue
            if self._canonical_resume_section(section.get("title") or section.get("type")) != "experience":
                continue
            return len(self._as_list(section.get("content")))
        return 0

    def _document_is_better(
        self,
        candidate_doc: dict[str, Any],
        candidate_issues: list[str],
        current_doc: dict[str, Any],
        current_issues: list[str],
        candidate_data: dict[str, Any],
    ) -> bool:
        """Prefer grounded output: fewer critical issues, then completeness, then fewer issues."""
        cand_critical = self._has_critical_quality_issues(candidate_issues)
        curr_critical = self._has_critical_quality_issues(current_issues)
        if cand_critical != curr_critical:
            return not cand_critical

        src_count = len(self._as_list(candidate_data.get("experience")))
        cand_roles = self._experience_role_count(candidate_doc)
        curr_roles = self._experience_role_count(current_doc)
        if src_count:
            cand_gap = max(0, src_count - cand_roles)
            curr_gap = max(0, src_count - curr_roles)
            if cand_gap != curr_gap:
                return cand_gap < curr_gap
        if len(candidate_issues) != len(current_issues):
            return len(candidate_issues) < len(current_issues)
        return cand_roles >= curr_roles

    def _enforce_document_reliability(
        self,
        document: dict[str, Any],
        candidate_data: dict[str, Any],
        metadata: dict[str, Any],
    ) -> dict[str, Any]:
        """Deterministically clean and stabilize resume output before render."""
        normalized = normalize_resume_document(document, candidate_data)
        normalized["client_name"] = candidate_data.get("client_name", "")

        sections = normalized.get("sections") or []
        if not isinstance(sections, list):
            sections = []

        # Prefer client order but always keep data-backed fallback sections.
        requested_order = metadata.get("sections") or metadata.get("section_order") or []
        rank: dict[str, int] = {}
        if isinstance(requested_order, list):
            for idx, sec_name in enumerate(requested_order):
                canonical = self._canonical_resume_section(sec_name)
                if canonical and canonical not in rank:
                    rank[canonical] = idx

        default_rank = {
            "summary": 0,
            "skills": 1,
            "experience": 2,
            "projects": 3,
            "education": 4,
            "certifications": 5,
            "achievements": 6,
            "languages": 7,
        }

        section_titles = {
            "summary": "PROFESSIONAL SUMMARY:",
            "skills": "TECHNICAL SKILLS:",
            "experience": "PROFESSIONAL EXPERIENCE:",
            "projects": "PROJECTS:",
            "education": "EDUCATION:",
            "certifications": "CERTIFICATIONS:",
            "achievements": "ACHIEVEMENTS:",
            "languages": "LANGUAGES:",
        }

        # Prefer human labels detected from the uploaded template; never use path-like
        # keys such as "summary.text" / "experience.items" as visible titles.
        section_labels = (metadata or {}).get("section_labels") or {}
        field_mapping = (metadata or {}).get("field_mapping") or {}
        label_sources: list[tuple[str, Any]] = []
        if isinstance(section_labels, dict):
            label_sources.extend(section_labels.items())
        if isinstance(field_mapping, dict):
            label_sources.extend(field_mapping.items())
        for key, label in label_sources:
            canonical_key = self._canonical_resume_section(key)
            if canonical_key not in section_titles or not label:
                continue
            titled = str(label).strip()
            if not titled or "." in titled:
                continue
            # Skip contact field keys accidentally present in field_mapping.
            if canonical_key in {"header", "name", "email", "phone", "location"}:
                continue
            # Never allow body lines / long phrases as section titles.
            clean = titled[:-1].strip() if titled.endswith(":") else titled
            if len(clean) > 36 or len(clean.split()) > 4 or "," in clean or "|" in clean:
                continue
            if not self._looks_like_section_heading(clean):
                continue
            titled = clean.upper()
            section_titles[canonical_key] = titled if titled.endswith(":") else f"{titled}:"

        cleaned_sections: list[dict[str, Any]] = []
        seen_sections: set[str] = set()
        header_contact = " | ".join(normalized.get("header", {}).get("contact") or [])
        for section in sections:
            if not isinstance(section, dict):
                continue
            sec_type = str(section.get("type") or "").strip().lower()
            title = str(section.get("title") or "").strip()
            content = section.get("content")
            canonical = self._resolve_section_canonical(title, sec_type)
            if not self._is_known_resume_section(canonical):
                continue
            if canonical in seen_sections and canonical in {"summary", "skills", "experience", "education"}:
                continue

            # Always use validated canonical titles — never company/role names as headings.
            styled_title = section_titles.get(canonical) or f"{canonical.replace('_', ' ').upper()}:"
            if styled_title and not styled_title.endswith(":"):
                styled_title += ":"
            forced_type = self._canonical_section_type(canonical)

            if forced_type in {"experience", "education", "projects"}:
                items = self._as_list(content)
                cleaned_items = []
                for item in items:
                    if not isinstance(item, dict):
                        text = self._clean_inline_text(item)
                        if text and not self._is_resume_noise_line(text, header_contact):
                            cleaned_items.append({"title": text, "description": []})
                        continue
                    clone = dict(item)
                    desc = clone.get("description") or clone.get("details") or clone.get("responsibilities") or clone.get("achievements") or []
                    desc_list = self._repair_bullet_sentences(self._as_list(desc), header_contact)
                    if desc_list:
                        clone["description"] = desc_list
                    for key in ("company", "institution", "organization", "school", "university", "title", "role", "position", "degree"):
                        if clone.get(key):
                            clean_val = self._clean_inline_text(clone.get(key))
                            clone[key] = clean_val if not self._is_resume_noise_line(clean_val, header_contact) else ""
                    cleaned_items.append(clone)
                if forced_type == "experience":
                    cleaned_items = self._filter_grounded_experience(
                        cleaned_items,
                        candidate_data.get("experience") or [],
                        header_contact,
                    )
                elif forced_type == "education":
                    cleaned_items = self._filter_grounded_education(
                        cleaned_items,
                        candidate_data.get("education") or [],
                    )
                if cleaned_items:
                    cleaned_sections.append({"type": forced_type, "title": styled_title, "content": cleaned_items})
                    seen_sections.add(canonical)
                continue

            if forced_type == "skills":
                grouped = self._filter_grounded_skills(content, candidate_data)
                if grouped:
                    cleaned_sections.append({"type": forced_type, "title": styled_title, "content": grouped})
                    seen_sections.add(canonical)
                continue

            if forced_type == "text":
                text_value = self._normalize_summary_text(content)
                if text_value:
                    cleaned_sections.append({"type": forced_type, "title": styled_title, "content": text_value})
                    seen_sections.add(canonical)
                continue

            bullets = [self._clean_inline_text(v) for v in self._as_list(content)]
            bullets = [
                b
                for b in self._dedupe_text(bullets)
                if b and not self._is_resume_noise_line(b, header_contact)
            ]
            if bullets:
                cleaned_sections.append({"type": forced_type, "title": styled_title, "content": bullets})
                seen_sections.add(canonical)

        # Ensure key sections use valid candidate-backed data and are not randomly omitted.
        for canonical in ("summary", "skills", "experience", "projects", "education", "certifications", "achievements", "languages"):
            if canonical in seen_sections:
                continue
            fallback_section = self._fallback_section_from_candidate(canonical, candidate_data)
            if fallback_section:
                fallback_section["title"] = section_titles.get(canonical, fallback_section.get("title"))
                cleaned_sections.append(fallback_section)
                seen_sections.add(canonical)

        cleaned_sections.sort(
            key=lambda sec: (
                rank.get(self._canonical_resume_section(sec.get("title") or sec.get("type")), 999),
                default_rank.get(self._canonical_resume_section(sec.get("title") or sec.get("type")), 999),
            )
        )
        normalized["sections"] = cleaned_sections
        return normalized

    def _restore_missing_roles(
        self,
        generated_items: list[dict[str, Any]],
        source_items: list[Any],
        header_contact: str = "",
    ) -> list[dict[str, Any]]:
        """Backfill any experience roles the LLM dropped from candidate data."""
        restored = list(generated_items)
        present = {self._experience_identity(item) for item in restored}
        for item in self._as_list(source_items):
            if not isinstance(item, dict):
                continue
            key = self._experience_identity(item)
            if not key or key == "||" or key in present:
                continue
            clone = dict(item)
            desc = clone.get("description") or clone.get("details") or clone.get("responsibilities") or []
            desc_list = self._repair_bullet_sentences(self._as_list(desc), header_contact)
            if desc_list:
                clone["description"] = desc_list
            restored.append(clone)
            present.add(key)
        return restored

    def _canonical_section_type(self, canonical: str) -> str:
        if canonical == "summary":
            return "text"
        if canonical == "skills":
            return "skills"
        if canonical in {"experience", "education", "projects"}:
            return canonical
        return "bullets"

    def _fallback_section_from_candidate(
        self, canonical: str, candidate_data: dict[str, Any]
    ) -> dict[str, Any] | None:
        titles = {
            "summary": "PROFESSIONAL SUMMARY:",
            "skills": "TECHNICAL SKILLS:",
            "experience": "PROFESSIONAL EXPERIENCE:",
            "projects": "PROJECTS:",
            "education": "EDUCATION:",
            "certifications": "CERTIFICATIONS:",
            "achievements": "ACHIEVEMENTS:",
            "languages": "LANGUAGES:",
        }
        title = titles.get(canonical, canonical.upper() + ":")
        if canonical == "summary":
            summary = self._normalize_summary_text(candidate_data.get("summary"))
            if summary:
                return {"type": "text", "title": title, "content": summary}
            return None
        if canonical == "skills":
            grouped = candidate_data.get("skills_by_category")
            if not isinstance(grouped, dict) or not grouped:
                grouped = self._group_skills_for_resume(candidate_data.get("skills"))
            if grouped:
                return {"type": "skills", "title": title, "content": grouped}
            return None
        if canonical in {"experience", "education", "projects"}:
            items = self._as_list(candidate_data.get(canonical))
            if items:
                return {"type": canonical, "title": title, "content": items}
            return None
        if canonical in {"certifications", "achievements", "languages"}:
            values = [self._clean_inline_text(v) for v in self._as_list(candidate_data.get(canonical))]
            values = [v for v in self._dedupe_text(values) if v]
            if values:
                return {"type": "bullets", "title": title, "content": values}
            return None
        return None

    def _canonical_resume_section(self, section_name: Any) -> str:
        normalized = re.sub(r"[^a-z0-9 ]+", " ", str(section_name or "").lower()).strip()
        normalized = re.sub(r"\s+", " ", normalized)
        if not normalized:
            return ""
        # Word-aware matching avoids false hits (e.g. "network" containing "work").
        if re.search(r"\b(experience|employment|work history|work experience|professional experience)\b", normalized):
            return "experience"
        if re.search(r"\b(education|academic|qualification|academics)\b", normalized):
            return "education"
        if re.search(r"\b(skills?|competenc\w*|expertise|technical skills)\b", normalized):
            return "skills"
        if re.search(r"\b(projects?|selected projects|project experience)\b", normalized):
            return "projects"
        if re.search(r"\b(certifications?|certificates?|licenses?)\b", normalized):
            return "certifications"
        if re.search(r"\b(summary|objective|profile summary|professional summary|about me)\b", normalized):
            return "summary"
        if re.search(r"\b(achievements?|awards?|honors?)\b", normalized):
            return "achievements"
        if re.search(r"\b(languages?)\b", normalized) and "programming" not in normalized:
            return "languages"
        return normalized

    def _is_known_resume_section(self, canonical: str) -> bool:
        return canonical in {
            "summary",
            "skills",
            "experience",
            "education",
            "projects",
            "certifications",
            "achievements",
            "languages",
        }

    def _looks_like_section_heading(self, title: Any) -> bool:
        """True only for real section headings — not company, role, or person names."""
        raw = str(title or "").strip()
        if not raw:
            return False
        clean = raw[:-1].strip() if raw.endswith(":") else raw
        if not clean or len(clean) > 55:
            return False
        low = clean.lower()
        if "@" in low or "http" in low or re.search(r"\d{3}[-.\s]?\d{3}", low):
            return False
        # Reject obvious employer / role shaped titles.
        if re.search(r"\b(inc|llc|ltd|corp|pvt|gmbh|technologies|solutions|consulting)\b", low):
            return False
        if re.search(r"\b(senior|junior|lead|principal|engineer|developer|manager|analyst|architect|intern)\b", low):
            # Allow only if also contains a section keyword ("Professional Experience").
            if not re.search(
                r"\b(experience|education|skills?|summary|projects?|certifications?|achievements?|languages?)\b",
                low,
            ):
                return False
        canonical = self._canonical_resume_section(clean)
        if not self._is_known_resume_section(canonical):
            return False
        # Must contain a recognizable section keyword, not just map via weak alias.
        keyword_map = {
            "summary": ("summary", "objective", "profile"),
            "skills": ("skill", "competenc", "expertise"),
            "experience": ("experience", "employment", "work history"),
            "education": ("education", "academic", "qualification"),
            "projects": ("project",),
            "certifications": ("certif", "license"),
            "achievements": ("achievement", "award", "honor"),
            "languages": ("language",),
        }
        return any(token in low for token in keyword_map.get(canonical, ()))

    def _resolve_section_canonical(self, title: Any, sec_type: Any) -> str:
        """Prefer validated heading text; fall back to section type when title is a company/role."""
        title_canon = self._canonical_resume_section(title)
        type_canon = self._canonical_resume_section(sec_type)
        if self._looks_like_section_heading(title) and self._is_known_resume_section(title_canon):
            return title_canon
        if self._is_known_resume_section(type_canon):
            return type_canon
        if self._is_known_resume_section(title_canon):
            return title_canon
        return ""

    def _normalize_entity_name(self, value: Any) -> str:
        text = re.sub(r"[^a-z0-9 ]+", " ", str(value or "").lower())
        text = re.sub(r"\s+", " ", text).strip()
        for noise in (" inc", " llc", " ltd", " corp", " pvt", " limited", " company"):
            if text.endswith(noise):
                text = text[: -len(noise)].strip()
        return text

    def _entity_matches(self, left: Any, right: Any) -> bool:
        a = self._normalize_entity_name(left)
        b = self._normalize_entity_name(right)
        if not a or not b:
            return False
        if a == b:
            return True
        if a in b or b in a:
            return True
        a_tokens = {t for t in a.split() if len(t) > 2}
        b_tokens = {t for t in b.split() if len(t) > 2}
        if not a_tokens or not b_tokens:
            return False
        overlap = len(a_tokens & b_tokens) / max(1, min(len(a_tokens), len(b_tokens)))
        return overlap >= 0.6

    def _source_experience_names(self, candidate_data: dict[str, Any]) -> set[str]:
        names: set[str] = set()
        for item in self._as_list(candidate_data.get("experience")):
            if not isinstance(item, dict):
                continue
            for key in ("company", "organization", "employer", "client"):
                val = self._normalize_entity_name(item.get(key))
                if val:
                    names.add(val)
            for key in ("title", "role", "position", "job_title"):
                val = self._normalize_entity_name(item.get(key))
                if val:
                    names.add(val)
        return names

    def _source_education_names(self, candidate_data: dict[str, Any]) -> set[str]:
        names: set[str] = set()
        for item in self._as_list(candidate_data.get("education")):
            if not isinstance(item, dict):
                continue
            for key in ("institution", "school", "university", "college", "organization"):
                val = self._normalize_entity_name(item.get(key))
                if val:
                    names.add(val)
            for key in ("degree", "field", "major"):
                val = self._normalize_entity_name(item.get(key))
                if val:
                    names.add(val)
        return names

    def _source_skill_names(self, candidate_data: dict[str, Any]) -> set[str]:
        return {self._normalize_entity_name(s) for s in self._clean_skill_list(candidate_data.get("skills")) if self._normalize_entity_name(s)}

    def _filter_grounded_experience(
        self,
        generated_items: list[dict[str, Any]],
        source_items: list[Any],
        header_contact: str = "",
    ) -> list[dict[str, Any]]:
        """Keep only experience entries grounded in candidate data; drop invented employers/roles."""
        source_list = [item for item in self._as_list(source_items) if isinstance(item, dict)]
        if not source_list:
            return []
        source_ids = {self._experience_identity(item) for item in source_list}
        grounded: list[dict[str, Any]] = []
        for item in generated_items:
            if not isinstance(item, dict):
                continue
            identity = self._experience_identity(item)
            if identity in source_ids and identity not in {"", "||"}:
                grounded.append(item)
                continue
            company = item.get("company") or item.get("organization") or item.get("employer") or ""
            title = item.get("title") or item.get("role") or item.get("position") or ""
            matched = False
            for src in source_list:
                src_company = src.get("company") or src.get("organization") or src.get("employer") or ""
                src_title = src.get("title") or src.get("role") or src.get("position") or ""
                if self._entity_matches(company, src_company) or (
                    self._entity_matches(title, src_title) and self._entity_matches(company, src_company)
                ):
                    matched = True
                    break
                if self._entity_matches(title, src_title) and not company:
                    matched = True
                    break
            if matched:
                grounded.append(item)
        restored = self._restore_missing_roles(grounded, source_list, header_contact)
        # Collapse near-duplicates created when company was temporarily blanked.
        return self._dedupe_experience_items(restored)

    def _dedupe_experience_items(self, items: list[dict[str, Any]]) -> list[dict[str, Any]]:
        """Keep one row per company+duration, preferring richer title/bullets."""
        out: list[dict[str, Any]] = []
        index: dict[str, int] = {}
        for item in items:
            if not isinstance(item, dict):
                continue
            company = str(item.get("company") or item.get("organization") or "").strip()
            duration = str(item.get("duration") or item.get("dates") or "").strip()
            key = re.sub(
                r"[^a-z0-9]+",
                "",
                f"{company}|{duration}".lower(),
            )
            if not key:
                # Fall back to full identity when company is empty.
                key = self._experience_identity(item)
            if not key or key == "||":
                out.append(item)
                continue
            if key in index:
                prev = out[index[key]]
                prev_bullets = prev.get("description") if isinstance(prev.get("description"), list) else []
                new_bullets = item.get("description") if isinstance(item.get("description"), list) else []
                prev_company = str(prev.get("company") or "").strip()
                new_company = str(item.get("company") or "").strip()
                richer = False
                if len(new_bullets) > len(prev_bullets):
                    richer = True
                if new_company and not prev_company:
                    richer = True
                if richer:
                    # Preserve a good title if the richer row lacks one.
                    if not item.get("title") and prev.get("title"):
                        item = {**item, "title": prev.get("title")}
                    if not new_company and prev_company:
                        item = {**item, "company": prev_company}
                    out[index[key]] = item
                elif item.get("title") and not prev.get("title"):
                    prev["title"] = item.get("title")
                elif new_company and not prev_company:
                    prev["company"] = new_company
                continue
            index[key] = len(out)
            out.append(item)
        return out

    def _filter_grounded_education(
        self,
        generated_items: list[dict[str, Any]],
        source_items: list[Any],
    ) -> list[dict[str, Any]]:
        source_list = [item for item in self._as_list(source_items) if isinstance(item, dict)]
        if not source_list:
            return []
        grounded: list[dict[str, Any]] = []
        for item in generated_items:
            if not isinstance(item, dict):
                continue
            institution = (
                item.get("institution")
                or item.get("school")
                or item.get("university")
                or item.get("college")
                or item.get("organization")
                or ""
            )
            degree = item.get("degree") or item.get("field") or ""
            for src in source_list:
                src_inst = (
                    src.get("institution")
                    or src.get("school")
                    or src.get("university")
                    or src.get("college")
                    or src.get("organization")
                    or ""
                )
                src_degree = src.get("degree") or src.get("field") or ""
                if self._entity_matches(institution, src_inst) or (
                    self._entity_matches(degree, src_degree) and institution
                ):
                    grounded.append(item)
                    break
        if not grounded:
            return [dict(item) for item in source_list]
        # Backfill any missing source schools.
        present = {
            self._normalize_entity_name(
                item.get("institution") or item.get("school") or item.get("university") or ""
            )
            for item in grounded
        }
        for src in source_list:
            key = self._normalize_entity_name(
                src.get("institution") or src.get("school") or src.get("university") or ""
            )
            if key and key not in present:
                grounded.append(dict(src))
                present.add(key)
        return grounded

    def _filter_grounded_skills(self, content: Any, candidate_data: dict[str, Any]) -> dict[str, list[str]]:
        """Keep skills that appear in source data; drop invented skill names."""
        source = self._source_skill_names(candidate_data)
        grouped = self._clean_skill_groups(content)
        if not grouped:
            grouped = self._group_skills_for_resume(self._clean_skill_list(content))
        if not source:
            return grouped or self._group_skills_for_resume(candidate_data.get("skills"))

        cleaned: dict[str, list[str]] = {}
        for category, values in (grouped or {}).items():
            kept: list[str] = []
            for skill in self._clean_skill_list(values):
                norm = self._normalize_entity_name(skill)
                if not norm:
                    continue
                if any(self._entity_matches(norm, src) for src in source):
                    kept.append(skill)
            if kept:
                cleaned[category] = kept
        if not cleaned:
            return self._group_skills_for_resume(candidate_data.get("skills"))
        return self._compact_skill_groups(cleaned)

    def _normalize_summary_text(self, value: Any) -> str:
        text = str(value or "")
        if not text.strip():
            return ""
        parts = [self._clean_inline_text(p) for p in re.split(r"[•\n]+", text)]
        parts = [p for p in parts if p and not self._is_resume_noise_line(p)]
        parts = self._dedupe_text(parts)
        if not parts:
            return ""
        # Preserve substantive summaries; polish happens optionally, not by hard clipping.
        joined = " ".join(parts)
        joined = re.sub(r"\s+", " ", joined).strip()
        return joined

    def _repair_bullet_sentences(self, values: list[Any], header_contact: str = "") -> list[str]:
        """Repair parser-split bullet lines into clean sentence bullets."""
        from app.services.structured_resume_store import expand_to_bullets

        # First expand paragraph streams into discrete bullets.
        expanded = expand_to_bullets(self._as_list(values))
        repaired: list[str] = []
        for raw in expanded:
            text = self._clean_inline_text(raw)
            if not text:
                continue
            if self._is_resume_noise_line(text, header_contact):
                continue
            # Drop obvious split artifacts.
            if self._looks_like_artifact_fragment(text):
                continue
            if repaired and self._should_merge_with_previous(text, repaired[-1]):
                repaired[-1] = f"{repaired[-1]} {text}".strip()
                continue
            repaired.append(text)

        final_items: list[str] = []
        for item in self._dedupe_text(repaired):
            item = re.sub(r"\s+\|\s+", " ", item).strip()
            item = re.sub(r"\s{2,}", " ", item).strip()
            if len(item) < 8:
                continue
            final_items.append(item)
        return final_items

    def _looks_like_artifact_fragment(self, text: str) -> bool:
        low = text.lower()
        if low in {"project", "impact", "technology", "technologies"}:
            return True
        if low.startswith("|") or low.endswith("|"):
            return True
        if " | " in low and len(low.split()) <= 4:
            return True
        return False

    def _should_merge_with_previous(self, text: str, previous: str = "") -> bool:
        """Merge only true wrap continuations — never glue separate achievements."""
        low = text.lower()
        prev = (previous or "").rstrip()
        # Previous already a complete sentence → keep as separate bullet.
        if prev.endswith((".", "!", "?")) and text[:1].isupper():
            return False
        if re.match(
            r"^(architected|developed|designed|implemented|built|led|created|improved|"
            r"optimized|deployed|integrated|managed|delivered|engineered)\b",
            low,
        ):
            return False
        if len(text.split()) <= 4 and not text[:1].isupper():
            return True
        if low.startswith(
            ("and ", "or ", "with ", "for ", "to ", "in ", "on ", "by ", "where ", "while ", "through ", "before ", "after ")
        ):
            return True
        if re.match(r"^[a-z].*", text) and len(text.split()) <= 12:
            return True
        return False

    def _is_resume_noise_line(self, line: str, header_contact: str = "") -> bool:
        low = str(line or "").strip().lower()
        if not low:
            return True
        if header_contact and low in header_contact.lower():
            return True
        # Only strip Aptino *template branding*, never a real employer named Aptino.
        noise_patterns = [
            r"^aptino\s+inc\.?\b",
            r"^aptino\s+corp\.?\b",
            r"^tushar\s+chouhan\b",
            r"^info@",
            r"^www\.",
            r"^\d{2,}\s+west\s+las\s+colinas\b",
            r"generated on",
            r"^email[:\s]",
            r"^address[:\s]",
            r"^phone[:\s]",
            r"^contact[:\s]",
            r"^https?://\S+$",
        ]
        if any(re.search(pattern, low) for pattern in noise_patterns):
            return True
        # Remove parser artifacts and truncated split markers.
        if " | " in low and ("project" in low or "impact" in low):
            return True
        if low.count("•") >= 2:
            return True
        return False

    def _resume_quality_feedback(
        self, candidate_data: dict[str, Any], document: dict[str, Any]
    ) -> list[str]:
        """Return a list of actionable issues; empty means OK."""
        issues: list[str] = []

        # Candidate content counts
        src_exp = candidate_data.get("experience") or []
        src_proj = candidate_data.get("projects") or []
        src_certs = candidate_data.get("certifications") or []
        src_skills = candidate_data.get("skills") or []

        # Document sections
        sections = document.get("sections") or []
        if not isinstance(sections, list) or not sections:
            return ["No sections were generated."]

        def _find_section(canonical: str) -> dict[str, Any] | None:
            for sec in sections:
                if not isinstance(sec, dict):
                    continue
                sec_type = str(sec.get("type") or "").lower().strip()
                title = str(sec.get("title") or "").lower().strip()
                resolved = self._resolve_section_canonical(sec.get("title"), sec.get("type"))
                if resolved == canonical or canonical in sec_type or canonical in title:
                    return sec
            return None

        exp_sec = _find_section("experience")
        proj_sec = _find_section("project")
        skills_sec = _find_section("skill")
        cert_sec = _find_section("cert")
        edu_sec = _find_section("education")

        # Experience completeness — require nearly all roles, even for short careers.
        if isinstance(src_exp, list) and src_exp:
            exp_items = exp_sec.get("content") if isinstance(exp_sec, dict) else None
            min_required = len(src_exp) if len(src_exp) <= 3 else max(len(src_exp) - 1, int(len(src_exp) * 0.9))
            if not isinstance(exp_items, list) or len(exp_items) < min_required:
                issues.append(
                    f"Experience section is incomplete ({len(exp_items) if isinstance(exp_items, list) else 0}/{len(src_exp)} roles); "
                    "include EVERY role from candidate data with full bullet lists."
                )
            elif isinstance(exp_items, list):
                # Detect over-summarized roles (too few bullets vs source).
                src_by_key = {
                    self._experience_identity(item): item
                    for item in src_exp
                    if isinstance(item, dict)
                }
                for item in exp_items:
                    if not isinstance(item, dict):
                        continue
                    src_item = src_by_key.get(self._experience_identity(item))
                    if not src_item:
                        continue
                    src_bullets = self._as_list(
                        src_item.get("description")
                        or src_item.get("details")
                        or src_item.get("responsibilities")
                    )
                    out_bullets = self._as_list(item.get("description") or item.get("details"))
                    if len(src_bullets) >= 4 and len(out_bullets) < max(3, int(len(src_bullets) * 0.6)):
                        issues.append(
                            "Some experience roles lost bullets; restore achievement details from candidate data."
                        )
                        break

        # Projects completeness
        if isinstance(src_proj, list) and len(src_proj) >= 2:
            proj_items = proj_sec.get("content") if isinstance(proj_sec, dict) else None
            if not isinstance(proj_items, list) or len(proj_items) < max(1, int(len(src_proj) * 0.7)):
                issues.append("Projects section is incomplete; include all projects from candidate data.")

        # Skills density
        if src_skills and skills_sec:
            skills_content = skills_sec.get("content")
            source_skill_count = len(self._clean_skill_list(src_skills))
            generated_skill_count = self._skill_content_count(skills_content)
            if generated_skill_count < min(12, source_skill_count):
                issues.append("Skills section is too short; include a fuller skills list and group logically.")
            if source_skill_count >= 8 and not isinstance(skills_content, dict):
                issues.append("Skills section is not grouped; use clear categories such as Programming Languages, Frameworks, Databases, Cloud & DevOps, and Tools.")

        # Certifications presence
        if isinstance(src_certs, list) and src_certs and not cert_sec:
            issues.append("Certifications are present in candidate data but missing in the resume.")

        # Required section preservation for client-ready resumes
        if src_skills and not skills_sec:
            issues.append("Skills are present in candidate data but missing in the resume.")
        if candidate_data.get("education") and not edu_sec:
            issues.append("Education is present in candidate data but missing in the resume.")

        # Format consistency: client style expects uppercase section titles with colon.
        invalid_titles = 0
        bad_heading_titles: list[str] = []
        for section in sections:
            if not isinstance(section, dict):
                continue
            title = str(section.get("title") or "").strip()
            if not title:
                continue
            normalized = title[:-1] if title.endswith(":") else title
            if normalized and normalized != normalized.upper():
                invalid_titles += 1
            sec_type = str(section.get("type") or "").strip()
            canonical = self._resolve_section_canonical(title, sec_type)
            if not self._looks_like_section_heading(title) and self._is_known_resume_section(canonical):
                # Title may still be wrong even if type is correct (company/role used as heading).
                if not re.search(
                    r"\b(summary|experience|education|skills?|projects?|certifications?|achievements?|languages?)\b",
                    title.lower(),
                ):
                    bad_heading_titles.append(title)
        if invalid_titles >= 2:
            issues.append("Section titles are not consistently formatted in client style (UPPERCASE with colon).")
        if bad_heading_titles:
            issues.append(
                "Invalid section title(s) look like company/role names instead of section headings "
                f"({', '.join(bad_heading_titles[:3])}); use PROFESSIONAL EXPERIENCE / EDUCATION / etc."
            )

        # Hallucination / grounding checks for experience employers.
        source_exp_names = self._source_experience_names(candidate_data)
        if source_exp_names and isinstance(exp_sec, dict):
            invented = []
            for item in self._as_list(exp_sec.get("content")):
                if not isinstance(item, dict):
                    continue
                company = item.get("company") or item.get("organization") or item.get("employer") or ""
                title = item.get("title") or item.get("role") or item.get("position") or ""
                company_norm = self._normalize_entity_name(company)
                title_norm = self._normalize_entity_name(title)
                company_ok = (not company_norm) or any(self._entity_matches(company_norm, src) for src in source_exp_names)
                title_ok = (not title_norm) or any(self._entity_matches(title_norm, src) for src in source_exp_names)
                if company_norm and not company_ok:
                    invented.append(str(company))
                elif title_norm and not title_ok and not company_ok:
                    invented.append(str(title))
            if invented:
                issues.append(
                    "Hallucinated experience entries detected "
                    f"({', '.join(invented[:3])}); use only employers/roles from candidate data."
                )

        # Hallucinated schools
        source_edu_names = self._source_education_names(candidate_data)
        if source_edu_names and isinstance(edu_sec, dict):
            invented_schools = []
            for item in self._as_list(edu_sec.get("content")):
                if not isinstance(item, dict):
                    continue
                institution = (
                    item.get("institution")
                    or item.get("school")
                    or item.get("university")
                    or item.get("college")
                    or ""
                )
                inst_norm = self._normalize_entity_name(institution)
                if inst_norm and not any(self._entity_matches(inst_norm, src) for src in source_edu_names):
                    invented_schools.append(str(institution))
            if invented_schools:
                issues.append(
                    "Hallucinated education institutions detected "
                    f"({', '.join(invented_schools[:3])}); use only schools from candidate data."
                )

        # Cross-section contamination
        if isinstance(skills_sec, dict):
            skill_blob = json.dumps(skills_sec.get("content") or "", ensure_ascii=True).lower()
            if re.search(r"\b(responsible for|led a team|managed project|years of experience)\b", skill_blob):
                issues.append(
                    "Section mix: Skills section contains experience-style sentences; keep skills as atomic skill names only."
                )
        if isinstance(exp_sec, dict):
            for item in self._as_list(exp_sec.get("content")):
                if not isinstance(item, dict):
                    continue
                # Company field accidentally set to a section heading.
                company = str(item.get("company") or "")
                if self._looks_like_section_heading(company):
                    issues.append(
                        "Section mix: Experience company field contains a section heading; keep employer names only."
                    )
                    break
                # Title field used as section name.
                role_title = str(item.get("title") or item.get("role") or "")
                if role_title and self._looks_like_section_heading(role_title) and "experience" in role_title.lower():
                    issues.append(
                        "Invalid section/role mix: experience role title looks like a section heading."
                    )
                    break
        if isinstance(edu_sec, dict):
            for item in self._as_list(edu_sec.get("content")):
                text_blob = json.dumps(item, ensure_ascii=True).lower() if isinstance(item, dict) else str(item).lower()
                if re.search(r"\b(developed|implemented|deployed|optimized microservices)\b", text_blob):
                    issues.append(
                        "Section mix: Education contains experience-style work bullets; keep academic details only."
                    )
                    break

        # Bullet quality: check for very short bullets (signals over-summarization)
        if isinstance(exp_sec, dict):
            exp_items = exp_sec.get("content") or []
            if isinstance(exp_items, list):
                short_bullets = 0
                total_bullets = 0
                noisy_bullets = 0
                artifact_bullets = 0
                for item in exp_items:
                    if not isinstance(item, dict):
                        continue
                    bullets = item.get("description") or item.get("details") or []
                    if isinstance(bullets, list):
                        for b in bullets:
                            total_bullets += 1
                            bullet_text = str(b).strip()
                            if len(bullet_text) < 45:
                                short_bullets += 1
                            if self._is_resume_noise_line(bullet_text):
                                noisy_bullets += 1
                            if self._looks_like_artifact_fragment(bullet_text) or " | " in bullet_text:
                                artifact_bullets += 1
                if total_bullets >= 10 and (short_bullets / max(1, total_bullets)) > 0.6:
                    issues.append(
                        "Experience bullets are too short/generic; rewrite with Action + Context + Result detail."
                    )
                if noisy_bullets >= 2:
                    issues.append(
                        "Experience includes contact/address/noise lines; keep only role responsibilities and impact."
                    )
                if artifact_bullets >= 2:
                    issues.append(
                        "Experience contains broken split fragments; merge into complete sentence bullets."
                    )

        # Detect jammed/mashed text (missing spaces between words).
        mashed_samples: list[str] = []
        summary_sec = _find_section("summary")
        if isinstance(summary_sec, dict) and self._text_looks_mashed(str(summary_sec.get("content") or "")):
            mashed_samples.append("summary")
        if isinstance(exp_sec, dict):
            for item in self._as_list(exp_sec.get("content")):
                if not isinstance(item, dict):
                    continue
                for b in self._as_list(item.get("description") or item.get("details")):
                    if self._text_looks_mashed(str(b)):
                        mashed_samples.append("experience")
                        break
                if "experience" in mashed_samples:
                    break
        if mashed_samples:
            issues.append(
                "Mashed/jammed text detected (missing spaces between words) in "
                f"{', '.join(mashed_samples)}; restore normal word spacing and professional sentences."
            )

        # Paragraph-stream detection: experience should be bullet-wise, not walls of text.
        if isinstance(exp_sec, dict):
            paragraph_roles = 0
            for item in self._as_list(exp_sec.get("content")):
                if not isinstance(item, dict):
                    continue
                bullets = self._as_list(item.get("description") or item.get("details"))
                if not bullets:
                    continue
                long_para = sum(1 for b in bullets if len(str(b)) > 350 and str(b).count(". ") >= 2)
                if len(bullets) <= 2 and long_para:
                    paragraph_roles += 1
            if paragraph_roles:
                issues.append(
                    "Experience uses paragraph streams instead of discrete bullets; "
                    "split into Action + Context + Result bullet points."
                )

        return issues

    def _text_looks_mashed(self, text: str) -> bool:
        """True when letters are jammed together with almost no spaces."""
        sample = re.sub(r"\s+", " ", str(text or "").strip())
        if len(sample) < 40:
            return False
        letters = sum(1 for c in sample if c.isalpha())
        if letters < 40:
            return False
        spaces = sample.count(" ")
        if spaces / max(letters, 1) < 0.08:
            return True
        # Very long alphabetic tokens are a strong mashed-text signal.
        long_tokens = re.findall(r"[A-Za-z]{22,}", sample)
        return len(long_tokens) >= 2

    def _has_critical_quality_issues(self, issues: list[str]) -> bool:
        critical_markers = (
            "hallucinated",
            "invalid section title",
            "section mix",
            "invented",
            "incomplete",
            "missing in the resume",
            "no sections",
            "mashed/jammed",
            "environment/tech lines",
            "paragraph streams",
            "fewer than 2 bullets",
            "education is missing",
            "skills section is too thin",
            "non-canonical section title",
            "truncated email",
            "valid email",
            "candidate name is missing",
            "certifications were mashed",
            "[critical]",
        )
        for issue in issues or []:
            low = str(issue).lower()
            if any(marker in low for marker in critical_markers):
                return True
        return False

    @traceable(run_type="llm", tags=["resume", "llm-compose", "client-ready"])
    def _compose_document_with_llm(
        self,
        candidate_data: dict[str, Any],
        metadata: dict[str, Any],
        baseline_document: dict[str, Any],
    ) -> dict[str, Any] | None:
        """Compose the resume JSON directly with the LLM using candidate facts."""
        compressed = self._compress_candidate_data(candidate_data)
        candidate_payload = json.dumps(compressed, ensure_ascii=True)[:70000]
        metadata_payload = json.dumps(metadata, ensure_ascii=True)[:9000]
        baseline_payload = json.dumps(baseline_document, ensure_ascii=True)[:45000]
        format_preview = str(metadata.get("preview_text") or "")[:3500]

        prompts = [
            (
                CLIENT_RESUME_GENERATOR_SYSTEM,
                CLIENT_RESUME_GENERATOR_PROMPT.format(
                    format_metadata=metadata_payload,
                    format_preview=format_preview,
                    candidate_data=candidate_payload,
                    baseline_resume=baseline_payload,
                ),
            ),
            (
                CLIENT_RESUME_GENERATOR_SYSTEM_COMPACT,
                CLIENT_RESUME_GENERATOR_PROMPT_COMPACT.format(
                    format_metadata=metadata_payload,
                    candidate_data=candidate_payload,
                    baseline_resume=baseline_payload,
                ),
            ),
        ]

        for attempt_index, (system_prompt, user_prompt) in enumerate(prompts, start=1):
            try:
                llm_result = llm_call_json_with_metrics(
                    system_prompt,
                    user_prompt,
                    validate=lambda data: self._validate_resume_document_json(data, candidate_data),
                    repair_attempts=1,
                    validation_attempts=1,
                    max_tokens=max(settings.RESUME_GENERATION_MAX_TOKENS, settings.LLM_MAX_TOKENS),
                )
                normalized = normalize_resume_document(llm_result.data, candidate_data)
                if self._document_is_substantially_thinner(normalized, baseline_document, candidate_data):
                    logger.warning(
                        "resume_llm_compose_rejected_thin",
                        attempt=attempt_index,
                        truncated=llm_result.metrics.get("truncated_output", False),
                        output_tokens_est=llm_result.metrics.get("output_tokens_est", 0),
                    )
                    continue
                logger.info(
                    "resume_llm_compose_complete",
                    attempt=attempt_index,
                    retry_count=llm_result.metrics.get("retry_count", 0),
                    json_repair_count=llm_result.metrics.get("json_repair_count", 0),
                    validation_retry_count=llm_result.metrics.get("validation_retry_count", 0),
                    input_tokens_est=llm_result.metrics.get("input_tokens_est", 0),
                    output_tokens_est=llm_result.metrics.get("output_tokens_est", 0),
                    truncated_output=llm_result.metrics.get("truncated_output", False),
                )
                return normalized
            except Exception as exc:
                logger.warning("resume_llm_compose_failed", attempt=attempt_index, error=str(exc))

        return None

    def _validate_resume_document_json(
        self,
        data: dict[str, Any],
        candidate_data: dict[str, Any],
    ) -> list[str]:
        """Validation feedback for LLM repair; deterministic guardrails do final cleanup."""
        errors: list[str] = []
        if not isinstance(data, dict):
            return ["Top-level response must be a JSON object."]

        sections = data.get("sections")
        if not isinstance(sections, list) or not sections:
            errors.append("Resume JSON must include a non-empty sections list.")
            return errors

        header = data.get("header")
        if not isinstance(header, dict):
            errors.append("Resume JSON must include a header object.")

        present = {
            self._canonical_resume_section((section or {}).get("title") or (section or {}).get("type"))
            for section in sections
            if isinstance(section, dict)
        }
        if candidate_data.get("skills") and "skills" not in present:
            errors.append("Candidate has skills; include a TECHNICAL SKILLS section.")
        if candidate_data.get("education") and "education" not in present:
            errors.append("Candidate has education; include an EDUCATION section.")
        if candidate_data.get("experience") and "experience" not in present:
            errors.append("Candidate has experience; include an EXPERIENCE section.")

        skills_section = next(
            (
                section
                for section in sections
                if isinstance(section, dict)
                and self._canonical_resume_section(section.get("title") or section.get("type")) == "skills"
            ),
            None,
        )
        if skills_section:
            skill_count = self._skill_content_count(skills_section.get("content"))
            source_count = len(self._clean_skill_list(candidate_data.get("skills")))
            if source_count >= 8 and skill_count < min(source_count, 12):
                errors.append("Technical skills section is too sparse; include the meaningful source skills.")
            if source_count >= 8 and not isinstance(skills_section.get("content"), dict):
                errors.append("Technical skills should be a categorized JSON object, not one flat paragraph.")

        for section in sections:
            if not isinstance(section, dict):
                continue
            title = str(section.get("title") or "").strip()
            sec_type = str(section.get("type") or "").strip()
            if title and not self._looks_like_section_heading(title):
                # Allow type-backed sections only when title is empty-ish; otherwise reject bad headings.
                if not re.search(
                    r"\b(summary|experience|education|skills?|projects?|certifications?|achievements?|languages?)\b",
                    title.lower(),
                ):
                    errors.append(
                        f"Invalid section title '{title}'. Use canonical headings like PROFESSIONAL EXPERIENCE, not company/role names."
                    )

            canonical = self._resolve_section_canonical(title, sec_type)
            content = section.get("content")
            if canonical == "experience" and isinstance(content, list):
                source_names = self._source_experience_names(candidate_data)
                if source_names:
                    for item in content:
                        if not isinstance(item, dict):
                            continue
                        company = item.get("company") or item.get("organization") or item.get("employer") or ""
                        if company and not any(self._entity_matches(company, src) for src in source_names):
                            errors.append(
                                f"Hallucinated employer '{company}' is not in candidate data. Use only real employers."
                            )
                            break
            if canonical == "skills":
                blob = json.dumps(content or "", ensure_ascii=True).lower()
                if re.search(r"\b(responsible for|led a team|managed project)\b", blob):
                    errors.append("Section mix: skills content must not include experience sentences.")

        return errors

    @traceable(run_type="chain", tags=["resume", "llm-polish", "quality-gate"])
    def _rewrite_document(
        self,
        draft_document: dict[str, Any],
        candidate_data: dict[str, Any],
        metadata: dict[str, Any],
        feedback: list[str],
    ) -> dict[str, Any] | None:
        """One-pass rewrite using the rewriter prompt to improve completeness."""
        try:
            compressed = self._compress_candidate_data(candidate_data)
            candidate_payload = json.dumps(compressed, ensure_ascii=True)[:50000]
            metadata_payload = json.dumps(metadata, ensure_ascii=True)[:6000]
            draft_payload = json.dumps(draft_document, ensure_ascii=True)[:45000]
            review_feedback = "\n".join(f"- {item}" for item in feedback)[:2500]

            llm_result = llm_call_json_with_metrics(
                CLIENT_RESUME_REWRITER_SYSTEM,
                CLIENT_RESUME_REWRITER_PROMPT.format(
                    format_metadata=metadata_payload,
                    candidate_data=candidate_payload,
                    draft_resume=draft_payload,
                    review_feedback=review_feedback,
                ),
                validate=lambda data: self._validate_resume_document_json(data, candidate_data),
                repair_attempts=1,
                validation_attempts=1,
                max_tokens=max(settings.RESUME_GENERATION_MAX_TOKENS, settings.LLM_MAX_TOKENS),
            )
            normalized = normalize_resume_document(llm_result.data, candidate_data)
            if self._document_is_substantially_thinner(normalized, draft_document, candidate_data):
                logger.warning(
                    "resume_rewrite_rejected_thin",
                    truncated=llm_result.metrics.get("truncated_output", False),
                    output_tokens_est=llm_result.metrics.get("output_tokens_est", 0),
                )
                return None
            logger.info(
                "resume_llm_polish_complete",
                retry_count=llm_result.metrics.get("retry_count", 0),
                json_repair_count=llm_result.metrics.get("json_repair_count", 0),
                validation_retry_count=llm_result.metrics.get("validation_retry_count", 0),
                input_tokens_est=llm_result.metrics.get("input_tokens_est", 0),
                output_tokens_est=llm_result.metrics.get("output_tokens_est", 0),
                truncated_output=llm_result.metrics.get("truncated_output", False),
            )
            return normalized
        except Exception as exc:
            logger.warning("resume_rewrite_failed", error=str(exc))
            return None

    def _metadata_for_generation(self, metadata: dict[str, Any]) -> dict[str, Any]:
        """Keep layout hints but remove sample resume content and large logo payloads."""
        allowed_keys = {
            "source_type",
            "template_id",
            "template_name",
            "sections",
            "section_order",
            "styling",
            "field_mapping",
            "section_labels",
            "layout",
            "company_header",
            "company_footer",
            "logo_count",
            "preview_text",
        }
        cleaned = {
            key: metadata[key]
            for key in allowed_keys
            if key in metadata and metadata[key] not in (None, "")
        }
        labels = cleaned.get("section_labels")
        if isinstance(labels, dict):
            defaults = {
                "summary": "PROFESSIONAL SUMMARY",
                "skills": "TECHNICAL SKILLS",
                "experience": "PROFESSIONAL EXPERIENCE",
                "projects": "PROJECTS",
                "education": "EDUCATION",
                "certifications": "CERTIFICATIONS",
                "achievements": "ACHIEVEMENTS",
                "languages": "LANGUAGES",
            }
            safe: dict[str, str] = {}
            for key, value in labels.items():
                canonical = self._canonical_resume_section(key)
                if canonical not in defaults:
                    continue
                text = str(value or "").strip()
                clean = text[:-1].strip() if text.endswith(":") else text
                if not clean or len(clean) > 36 or len(clean.split()) > 4:
                    continue
                if not self._looks_like_section_heading(clean):
                    continue
                safe[canonical] = clean.upper()
            cleaned["section_labels"] = safe or defaults
        return cleaned

    def _valid_logos(self, logos: Any) -> list[dict]:
        import base64

        valid: list[dict] = []
        if not isinstance(logos, list):
            return valid
        for logo in logos:
            if not isinstance(logo, dict):
                continue
            data = str(logo.get("data") or "")
            if not data.startswith("data:image") or "," not in data:
                continue
            try:
                payload = data.split(",", 1)[1]
                # Padding-tolerant decode — many DOCX-extracted payloads lack strict padding.
                decoded = base64.b64decode(payload + "===", validate=False)
            except Exception:
                continue
            if 200 <= len(decoded) <= 5_000_000:
                valid.append(logo)
        if not valid:
            return []

        def _rank(item: dict) -> tuple[int, int]:
            source = str(item.get("source") or "").lower()
            position = str(item.get("position") or "").lower()
            source_rank = 0
            if "docx_header" in source or "pymupdf_header" in source or "aptino" in source:
                source_rank = 5
            elif "footer" in source or "footer" in position:
                source_rank = 4
            elif "header" in source:
                source_rank = 3
            elif "top_body" in source or "top_table" in source:
                source_rank = 2
            elif "package" in source or "pdf2image" in source:
                source_rank = 1
            return (source_rank, len(str(item.get("data") or "")))

        header_candidates = [
            logo for logo in valid
            if "footer" not in str(logo.get("position") or "").lower()
            and "footer" not in str(logo.get("source") or "").lower()
        ]
        footer_candidates = [
            logo for logo in valid
            if "footer" in str(logo.get("position") or "").lower()
            or "footer" in str(logo.get("source") or "").lower()
        ]

        selected: list[dict] = []
        if header_candidates:
            selected.append(sorted(header_candidates, key=_rank, reverse=True)[0])
        elif valid:
            # Fallback: best overall image as header logo.
            selected.append(sorted(valid, key=_rank, reverse=True)[0])
        if footer_candidates:
            best_footer = sorted(footer_candidates, key=_rank, reverse=True)[0]
            if best_footer not in selected:
                selected.append(best_footer)
        return selected

    def _split_logos(self, logos: list[dict]) -> tuple[list[dict], list[dict]]:
        """Split validated logos into header brand mark vs footer stamp/sign."""
        header: list[dict] = []
        footer: list[dict] = []
        for logo in logos or []:
            position = str(logo.get("position") or "").lower()
            source = str(logo.get("source") or "").lower()
            if "footer" in position or "footer" in source:
                footer.append(logo)
            else:
                header.append(logo)
        return header, footer

    def _clean_skill_groups(self, content: Any) -> dict[str, list[str]]:
        if not isinstance(content, dict):
            return self._group_skills_for_resume(content)

        grouped: dict[str, list[str]] = {}
        for category, values in content.items():
            category_name = self._normalize_skill_category(category)
            skills = self._clean_skill_list(values)
            if category_name and skills:
                grouped.setdefault(category_name, [])
                grouped[category_name].extend(skills)

        return self._compact_skill_groups(grouped)

    def _group_skills_for_resume(self, source: Any) -> dict[str, list[str]]:
        """Create recruiter-friendly skill categories from atomic skill names."""
        if isinstance(source, dict):
            flattened: list[str] = []
            for values in source.values():
                flattened.extend(self._clean_skill_list(values))
            skills = self._clean_skill_list(flattened)
        else:
            skills = self._clean_skill_list(source)

        if not skills:
            return {}

        category_keywords: list[tuple[str, tuple[str, ...]]] = [
            (
                "Programming Languages",
                (
                    "python",
                    "java",
                    "javascript",
                    "typescript",
                    "c++",
                    "c#",
                    "go",
                    "golang",
                    "ruby",
                    "php",
                    "swift",
                    "kotlin",
                    "scala",
                    "r",
                    "matlab",
                    "bash",
                    "shell",
                    "powershell",
                ),
            ),
            (
                "Frontend",
                (
                    "react",
                    "angular",
                    "vue",
                    "next.js",
                    "nuxt",
                    "html",
                    "css",
                    "tailwind",
                    "bootstrap",
                    "redux",
                    "material ui",
                    "mui",
                    "responsive design",
                ),
            ),
            (
                "Backend & Frameworks",
                (
                    "spring",
                    "spring boot",
                    "django",
                    "flask",
                    "fastapi",
                    "express",
                    "node.js",
                    "nodejs",
                    ".net",
                    "hibernate",
                    "jpa",
                    "rest api",
                    "graphql",
                    "microservices",
                    "api",
                ),
            ),
            (
                "Databases",
                (
                    "sql",
                    "mysql",
                    "postgresql",
                    "postgres",
                    "mongodb",
                    "oracle",
                    "sqlite",
                    "redis",
                    "cassandra",
                    "dynamodb",
                    "snowflake",
                    "bigquery",
                    "elasticsearch",
                ),
            ),
            (
                "Cloud & DevOps",
                (
                    "aws",
                    "azure",
                    "gcp",
                    "google cloud",
                    "docker",
                    "kubernetes",
                    "k8s",
                    "terraform",
                    "ansible",
                    "jenkins",
                    "ci/cd",
                    "github actions",
                    "gitlab ci",
                    "devops",
                    "linux",
                    "unix",
                ),
            ),
            (
                "Data & AI",
                (
                    "machine learning",
                    "deep learning",
                    "nlp",
                    "computer vision",
                    "pandas",
                    "numpy",
                    "spark",
                    "hadoop",
                    "tensorflow",
                    "pytorch",
                    "scikit-learn",
                    "keras",
                    "power bi",
                    "tableau",
                    "analytics",
                ),
            ),
            (
                "Testing & QA",
                (
                    "selenium",
                    "cypress",
                    "playwright",
                    "jest",
                    "pytest",
                    "junit",
                    "testng",
                    "postman",
                    "unit testing",
                    "automation testing",
                    "manual testing",
                    "qa",
                ),
            ),
            (
                "Tools & Platforms",
                (
                    "git",
                    "github",
                    "gitlab",
                    "bitbucket",
                    "jira",
                    "confluence",
                    "swagger",
                    "figma",
                    "maven",
                    "gradle",
                    "npm",
                    "yarn",
                    "webpack",
                ),
            ),
            (
                "Architecture & Practices",
                (
                    "agile",
                    "scrum",
                    "kanban",
                    "system design",
                    "design patterns",
                    "architecture",
                    "distributed systems",
                    "oop",
                    "object oriented",
                    "clean code",
                    "code review",
                ),
            ),
            (
                "Business & Domain",
                (
                    "banking",
                    "healthcare",
                    "fintech",
                    "ecommerce",
                    "retail",
                    "crm",
                    "erp",
                    "salesforce",
                    "sap",
                    "servicenow",
                    "workday",
                ),
            ),
            (
                "Soft Skills",
                (
                    "leadership",
                    "communication",
                    "collaboration",
                    "stakeholder",
                    "mentoring",
                    "problem solving",
                    "team management",
                    "presentation",
                ),
            ),
        ]

        grouped: dict[str, list[str]] = {category: [] for category, _ in category_keywords}
        grouped["Additional Skills"] = []

        for skill in skills:
            low = skill.casefold()
            selected_category = ""
            for category, keywords in category_keywords:
                if any(self._skill_matches_keyword(low, keyword) for keyword in keywords):
                    selected_category = category
                    break
            grouped[selected_category or "Additional Skills"].append(skill)

        return self._compact_skill_groups(grouped)

    def _compact_skill_groups(self, grouped: dict[str, list[str]]) -> dict[str, list[str]]:
        compacted: dict[str, list[str]] = {}
        for category, values in grouped.items():
            skills = self._dedupe_text([self._normalize_skill_name(value) for value in values])
            skills = [skill for skill in skills if skill and not self._is_bad_skill_name(skill)]
            if skills:
                compacted[category] = skills
        return compacted

    def _clean_skill_list(self, source: Any) -> list[str]:
        raw_items = self._extract_skill_items(source)
        skills = [self._normalize_skill_name(item) for item in raw_items]
        skills = [skill for skill in skills if skill and not self._is_bad_skill_name(skill)]
        # No hard cap — large resumes can list 100+ skills across categories.
        return self._dedupe_text(skills)

    def _normalize_skill_name(self, value: Any) -> str:
        from app.services.tech_glossary import normalize_skill_token

        skill = self._clean_inline_text(value)
        if not skill:
            return ""
        skill = re.sub(
            r"^(technical skills?|core skills?|key skills?|skills?|tools?|technologies?)\s*[:\-]\s*",
            "",
            skill,
            flags=re.IGNORECASE,
        )
        skill = skill.strip(" .,:;|/")
        skill = re.sub(r"\s*/\s*", "/", skill)
        skill = re.sub(r"\s*&\s*", " & ", skill)
        skill = re.sub(r"\s+", " ", skill).strip()
        return normalize_skill_token(skill)

    def _normalize_skill_category(self, value: Any) -> str:
        category = self._clean_inline_text(value).strip(" :")
        low = category.casefold()
        # Exact-ish aliases first — never let "soft" match inside "software".
        if re.search(r"\bsoft\s*skills?\b", low):
            return "Soft Skills"
        if "software development" in low or re.search(r"\bsoftware\b", low):
            return "Software Development"
        if "generative" in low or "llm" in low:
            return "Data & AI"
        mapping = {
            "language": "Programming Languages",
            "programming": "Programming Languages",
            "frontend": "Frontend",
            "front end": "Frontend",
            "backend": "Backend & Frameworks",
            "back end": "Backend & Frameworks",
            "framework": "Backend & Frameworks",
            "database": "Databases",
            "cloud": "Cloud & DevOps",
            "devops": "Cloud & DevOps",
            "machine learning": "Data & AI",
            "testing": "Testing & QA",
            "qa": "Testing & QA",
            "tool": "Tools & Platforms",
            "platform": "Tools & Platforms",
            "architecture": "Architecture & Practices",
            "practice": "Architecture & Practices",
            "methodolog": "Architecture & Practices",
            "business": "Business & Domain",
            "domain": "Business & Domain",
            "additional": "Additional Skills",
            "other": "Additional Skills",
        }
        # Prefer longer markers first to avoid partial collisions.
        for marker, normalized in sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True):
            if marker in low:
                return normalized
        # Short tokens that are unsafe as substrings.
        if re.search(r"(?i)\b(data|ai)\b", low):
            return "Data & AI"
        if category and len(category) <= 36:
            return category
        return "Additional Skills"

    def _is_bad_skill_name(self, skill: str) -> bool:
        low = skill.casefold().strip()
        if not low or low in {"n/a", "na", "none", "null", "skills"}:
            return True
        if "@" in low or "http://" in low or "https://" in low:
            return True
        if len(skill) > 70 or len(skill.split()) > 8:
            return True
        if low.endswith(":"):
            return True
        bad_starts = (
            "responsible for",
            "worked on",
            "experience with",
            "experienced in",
            "developed ",
            "implemented ",
            "managed ",
            "created ",
            "designed ",
            "built ",
            "led ",
        )
        if any(low.startswith(prefix) for prefix in bad_starts):
            return True
        sentence_verbs = (
            " developed ",
            " implemented ",
            " managed ",
            " created ",
            " designed ",
            " built ",
            " led ",
            " reduced ",
            " improved ",
            " collaborated ",
            " responsible ",
        )
        padded = f" {low} "
        return len(skill.split()) > 4 and any(verb in padded for verb in sentence_verbs)

    def _skill_matches_keyword(self, skill_lower: str, keyword: str) -> bool:
        keyword_lower = keyword.casefold()
        if not keyword_lower:
            return False
        if re.fullmatch(r"[a-z0-9+#./]+", keyword_lower):
            return bool(re.search(rf"(^|[^a-z0-9+#]){re.escape(keyword_lower)}([^a-z0-9+#]|$)", skill_lower))
        return keyword_lower in skill_lower

    def _skill_content_count(self, content: Any) -> int:
        if isinstance(content, dict):
            return sum(len(self._clean_skill_list(values)) for values in content.values())
        return len(self._clean_skill_list(content))

    def _extract_resume_text_from_path(self, candidate: Candidate) -> str:
        object_or_path = str(candidate.resume_path or "")
        if not object_or_path:
            return ""

        local_path = Path(object_or_path)
        downloaded_path: Path | None = None
        try:
            if local_path.exists():
                return extract_text_from_document(str(local_path))

            downloaded_path = Path(
                download_to_local_path(str(candidate.recruiter_id), object_or_path)
            )
            return extract_text_from_document(str(downloaded_path))
        finally:
            if downloaded_path:
                downloaded_path.unlink(missing_ok=True)

    def _extract_skill_items(self, source: Any) -> list[str]:
        if not source:
            return []
        if isinstance(source, dict):
            items: list[str] = []
            for key in ("items", "values", "skills", "matched", "not_matched"):
                if source.get(key):
                    items.extend(self._extract_skill_items(source[key]))
            if items:
                return items
            for value in source.values():
                items.extend(self._extract_skill_items(value))
            return items
        if isinstance(source, (list, tuple, set)):
            items = []
            for value in source:
                items.extend(self._extract_skill_items(value))
            return items
        if isinstance(source, str):
            from app.services.detailed_resume_parser import _split_skills_payload

            return [
                self._clean_inline_text(part)
                for part in _split_skills_payload(source)
                if self._clean_inline_text(part)
            ]
        return [self._clean_inline_text(source)]

    def _extract_skills_from_text(self, resume_text: str) -> list[str]:
        lines = self._section_lines(
            resume_text,
            ["skills", "technical skills", "core skills", "competencies", "technical expertise"],
        )
        skills: list[str] = []
        for line in lines:
            skills.extend(self._extract_skill_items(line))
        return [
            skill
            for skill in self._dedupe_text(skills)
            if 1 <= len(skill.split()) <= 8 and len(skill) <= 80
        ]

    def _extract_education_from_text(self, resume_text: str) -> list[dict[str, Any]]:
        lines = self._section_lines(
            resume_text,
            ["education", "academic background", "academics", "qualifications"],
        )
        if not lines:
            return []

        blocks: list[list[str]] = []
        current: list[str] = []
        for line in lines:
            clean = self._clean_inline_text(line)
            if not clean:
                continue
            if current and self._looks_like_degree(clean):
                blocks.append(current)
                current = [clean]
            else:
                current.append(clean)
        if current:
            blocks.append(current)

        education: list[dict[str, Any]] = []
        for block in blocks:
            degree = next((line for line in block if self._looks_like_degree(line)), block[0])
            year = next((line for line in block if re.search(r"\b(19|20)\d{2}\b", line)), "")
            institution = ""
            details: list[str] = []
            for line in block:
                if line in {degree, year}:
                    continue
                if not institution and not self._looks_like_detail(line):
                    institution = line
                else:
                    details.append(line)
            education.append(
                {
                    "degree": degree,
                    "institution": institution,
                    "year": year,
                    "details": self._dedupe_text(details),
                }
            )
        return education

    def _section_lines(self, text: str, heading_aliases: list[str]) -> list[str]:
        headings = [
            "summary",
            "professional summary",
            "profile",
            "objective",
            "skills",
            "technical skills",
            "core skills",
            "competencies",
            "technical expertise",
            "experience",
            "work experience",
            "professional experience",
            "employment",
            "education",
            "academic background",
            "academics",
            "qualifications",
            "projects",
            "certifications",
            "achievements",
            "languages",
        ]
        target_headings = {self._normalize_heading(alias) for alias in heading_aliases}
        all_headings = {self._normalize_heading(alias) for alias in headings}

        captured: list[str] = []
        in_section = False
        for raw_line in (text or "").splitlines():
            line = raw_line.strip()
            normalized = self._normalize_heading(line)
            if not line:
                if in_section:
                    captured.append("")
                continue
            if normalized in target_headings or any(
                normalized.startswith(f"{heading} ") for heading in target_headings
            ):
                in_section = True
                continue
            if in_section and (
                normalized in all_headings
                or any(normalized.startswith(f"{heading} ") for heading in all_headings)
            ):
                break
            if in_section:
                captured.append(line)
        return captured

    def _normalize_heading(self, value: str) -> str:
        return re.sub(r"[^a-z0-9 ]+", "", str(value).lower()).strip()

    def _looks_like_degree(self, line: str) -> bool:
        return bool(
            re.search(
                r"\b(bachelor|master|phd|doctor|b\.?tech|m\.?tech|b\.?e\.?|m\.?e\.?|"
                r"b\.?sc|m\.?sc|b\.?com|m\.?com|bca|mca|bba|bdes|mdes|mba|"
                r"degree|diploma)\b",
                line,
                flags=re.IGNORECASE,
            )
        )

    def _looks_like_detail(self, line: str) -> bool:
        return bool(
            re.search(
                r"\b(cgpa|gpa|percentage|coursework|honor|specialization)\b",
                line,
                re.IGNORECASE,
            )
        )

    def _clean_inline_text(self, value: Any) -> str:
        # Strip common bullet glyphs copied from source resumes to avoid double bullets in DOCX lists.
        from app.services.tech_glossary import restore_tech_names

        cleaned = re.sub(r"^[\-\*\u2022\u2023\u25AA\u25AB\u25CF\u25E6\u2043\u2219\uf0a7\uf0b7▪■●◦•\s]+", "", str(value or "")).strip()
        cleaned = re.sub(r"\s+", " ", cleaned)
        cleaned = restore_tech_names(cleaned)
        cleaned = re.sub(r"(?i)\b(AI/?ML)(Developer|Engineer|Architect|Specialist)\b", r"\1 \2", cleaned)
        return cleaned

    def _dedupe_text(self, values: list[str]) -> list[str]:
        seen: set[str] = set()
        result: list[str] = []
        for value in values:
            clean = self._clean_inline_text(value)
            if not clean:
                continue
            key = clean.casefold()
            if key in seen:
                continue
            seen.add(key)
            result.append(clean)
        return result

    def _candidate_to_data(self, candidate: Candidate) -> dict[str, Any]:
        """Legacy method - kept for backward compatibility."""
        extracted = candidate.extracted_data or {}
        return {
            "name": candidate.name or extracted.get("name"),
            "email": candidate.email or extracted.get("email"),
            "phone": candidate.phone or extracted.get("phone"),
            "location": candidate.location or extracted.get("location"),
            "job_applied": candidate.job_applied,
            "job_role": candidate.job_role,
            "client_name": candidate.client_name,
            "summary": extracted.get("summary", ""),
            "skills": candidate.skills_matched or extracted.get("skills", []),
            "experience": extracted.get("experience") or [],
            "education": extracted.get("education") or [],
            "projects": extracted.get("projects", []),
            "certifications": extracted.get("certifications", []),
            "extracted_data": extracted,
        }

    def _render_pdf(self, document: dict[str, Any], metadata: dict[str, Any] | None) -> bytes:
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Image
            from reportlab.platypus.frames import Frame
            from reportlab.platypus.doctemplate import PageTemplate
        except ImportError as exc:  # pragma: no cover - depends on runtime deps
            raise ResumeGenerationError("reportlab is required to generate client resumes") from exc

        styling = (metadata or {}).get("styling") or {}
        body_size = int(styling.get("font_size_body") or 11)
        header_size = int(styling.get("font_size_header") or 16)
        
        # Get client info for footer
        client_name = document.get("client_name") or ""
        from datetime import datetime
        footer_text = f"{client_name} - Generated on {datetime.now().strftime('%B %d, %Y')}" if client_name else f"Generated on {datetime.now().strftime('%B %d, %Y')}"
        
        # Logo URL from metadata
        logo_url = metadata.get("logo_url") if metadata else None

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=16 * mm,
            leftMargin=16 * mm,
            topMargin=20 * mm if logo_url else 14 * mm,
            bottomMargin=18 * mm,
        )
        
        # Define header/footer callback
        def header_footer(canvas, doc):
            canvas.saveState()
            
            # Add logo at top right if available
            if logo_url:
                try:
                    # Logo size: max 30mm width, 15mm height
                    logo = Image(logo_url, width=30*mm, height=15*mm)
                    logo.drawOn(canvas, doc.width + doc.leftMargin - 30*mm, doc.height + doc.topMargin - 10*mm)
                except Exception:
                    # If logo fails to load, skip it
                    pass
            
            # Add footer at bottom
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(colors.HexColor("#6b7280"))
            canvas.drawCentredString(
                doc.width / 2 + doc.leftMargin,
                10 * mm,
                footer_text
            )
            
            # Add line above footer
            canvas.setStrokeColor(colors.HexColor("#e5e7eb"))
            canvas.setLineWidth(0.5)
            canvas.line(doc.leftMargin, 14 * mm, doc.width + doc.leftMargin, 14 * mm)
            
            canvas.restoreState()
        
        # Add page template with header/footer
        frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id='normal')
        template = PageTemplate(id='resume_template', frames=frame, onPage=header_footer)
        doc.addPageTemplates([template])
        
        base_styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ClientResumeTitle",
            parent=base_styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=header_size + 4,
            leading=header_size + 8,
            textColor=colors.HexColor("#111827"),
            spaceAfter=4,
        )
        role_style = ParagraphStyle(
            "ClientResumeRole",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=body_size + 1,
            leading=body_size + 5,
            textColor=colors.HexColor("#374151"),
            spaceAfter=2,
        )
        contact_style = ParagraphStyle(
            "ClientResumeContact",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=max(9, body_size - 1),
            leading=body_size + 3,
            textColor=colors.HexColor("#4b5563"),
            spaceAfter=12,
        )
        section_style = ParagraphStyle(
            "ClientResumeSection",
            parent=base_styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=header_size,
            leading=header_size + 4,
            textColor=colors.HexColor("#111827"),
            borderColor=colors.HexColor("#d1d5db"),
            borderWidth=0,
            borderPadding=0,
            spaceBefore=10,
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            "ClientResumeBody",
            parent=base_styles["BodyText"],
            fontName="Helvetica",
            fontSize=body_size,
            leading=body_size + 4,
            textColor=colors.HexColor("#1f2937"),
            spaceAfter=5,
        )
        item_title_style = ParagraphStyle(
            "ClientResumeItemTitle",
            parent=body_style,
            fontName="Helvetica-Bold",
            spaceAfter=2,
        )
        bullet_style = ParagraphStyle(
            "ClientResumeBullet",
            parent=body_style,
            leftIndent=10,
            firstLineIndent=-7,
            spaceAfter=3,
        )

        story: list[Any] = []
        header = document.get("header") or {}
        story.append(Paragraph(escape(str(header.get("name") or "Candidate")), title_style))
        if header.get("role"):
            story.append(Paragraph(escape(str(header["role"])), role_style))
        if header.get("contact"):
            story.append(Paragraph(escape(" | ".join(header["contact"])), contact_style))

        for section in document.get("sections") or []:
            story.append(Paragraph(escape(str(section["title"])), section_style))
            self._append_section(story, section, body_style, item_title_style, bullet_style, Paragraph, Spacer)

        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    def _render_docx(
        self,
        document: dict[str, Any],
        metadata: dict[str, Any] | None,
        logos: list[dict],
        footer_logos: list[dict] | None = None,
    ) -> bytes:
        """Render resume as DOCX with Aptino-style single-column ATS layout."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.oxml.ns import qn
            from io import BytesIO
        except ImportError as exc:
            raise ResumeGenerationError("python-docx is required to generate DOCX resumes") from exc

        doc = Document()
        styling = (metadata or {}).get("styling") or {}
        body_size = float(styling.get("font_size_body") or 11)
        header_size = float(styling.get("font_size_header") or 12)
        name_size = float(styling.get("font_size_name") or 20)
        font_family = styling.get("font_family") or "Calibri"
        margin_inches = float(styling.get("margin_inches") or 0.65)

        style = doc.styles["Normal"]
        style.font.name = font_family
        style.font.size = Pt(body_size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)

        section = doc.sections[0]
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(max(0.6, margin_inches))
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

        usable_width = section.page_width - section.left_margin - section.right_margin
        header = document.get("header") or {}
        company_lines = (
            document.get("company_footer_lines")
            or document.get("company_header_lines")
            or []
        )
        if not company_lines:
            company_block = (metadata or {}).get("company_footer") or (metadata or {}).get("company_header") or {}
            if isinstance(company_block, dict):
                company_lines = list(company_block.get("lines") or [])

        logo_bytes = self._decode_logo_bytes(
            logos[0] if logos else None,
            strip_black_bg=is_aptino_template(metadata) or (
                bool(logos) and "aptino" in str((logos[0] or {}).get("source") or "").lower()
            ),
        )
        footer_stamp_bytes = self._decode_logo_bytes(
            (footer_logos or [None])[0] if footer_logos else None,
            strip_black_bg=False,
        )

        # Page header (every page): candidate name + position + logo. No company sign here.
        self._fill_docx_page_header(
            section.header,
            header,
            logo_bytes,
            usable_width,
            name_size,
            body_size,
            font_family,
        )

        # Page footer (every page): optional stamp image + company sign text.
        self._fill_docx_page_footer(
            section.footer,
            company_lines,
            font_family,
            stamp_bytes=footer_stamp_bytes,
        )

        # Contact on one line beneath the repeating header (body, first page flow).
        # Always put the name in the body so copy/paste, text extractors, and some
        # viewers show it even when page headers are hidden.
        display_name = str(header.get("name") or "Candidate").strip() or "Candidate"
        name_para = doc.add_paragraph()
        name_para.paragraph_format.space_before = Pt(0)
        name_para.paragraph_format.space_after = Pt(2)
        name_run = name_para.add_run(display_name)
        name_run.font.name = font_family
        name_run.font.size = Pt(max(name_size, body_size + 2))
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        if header.get("contact"):
            contact_para = doc.add_paragraph()
            contact_para.paragraph_format.space_before = Pt(2)
            contact_para.paragraph_format.space_after = Pt(6)
            contact_items = [str(c).strip() for c in header["contact"] if str(c).strip()]
            contact_run = contact_para.add_run("  |  ".join(contact_items))
            contact_run.font.name = font_family
            contact_run.font.size = Pt(max(9.5, body_size - 0.5))
            contact_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        self._add_horizontal_line(doc)

        for section_data in document.get("sections") or []:
            self._add_docx_section(doc, section_data, header_size, body_size, font_family, usable_width)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _clear_docx_container(self, container: Any) -> None:
        """Remove all paragraphs/tables from a header or footer container."""
        from docx.oxml.ns import qn

        element = container._element
        for child in list(element):
            if child.tag in {qn("w:p"), qn("w:tbl")}:
                element.remove(child)

    def _fill_docx_page_header(
        self,
        page_header: Any,
        header: dict,
        logo_bytes: bytes | None,
        usable_width: Any,
        name_size: float,
        body_size: float,
        font_family: str,
    ) -> None:
        """Put logo in the Word section header (repeats on every page).

        Candidate name is rendered once in the document body to avoid a
        double-name look (header + body) on page 1 when a client logo is present.
        """
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO

        self._clear_docx_container(page_header)

        if logo_bytes:
            table = page_header.add_table(1, 2, width=usable_width)
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = int(usable_width * 0.72)
            table.columns[1].width = int(usable_width * 0.28)

            left_cell = table.cell(0, 0)
            left_cell.paragraphs[0].clear()
            if header.get("role"):
                role_para = left_cell.paragraphs[0]
                role_para.paragraph_format.space_after = Pt(0)
                role_run = role_para.add_run(str(header["role"]))
                role_run.font.name = font_family
                role_run.font.size = Pt(body_size + 0.5)
                role_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            right_cell = table.cell(0, 1)
            right_cell.paragraphs[0].clear()
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            logo_run = right_para.add_run()
            logo_width = Inches(1.5)
            try:
                from PIL import Image as PILImage

                pil_img = PILImage.open(BytesIO(logo_bytes))
                img_width, img_height = pil_img.size
                aspect = img_width / img_height if img_height else 2
                if aspect > 3:
                    logo_width = Inches(1.7)
                elif aspect < 1.2:
                    logo_width = Inches(1.05)
            except Exception:
                pass
            try:
                logo_run.add_picture(BytesIO(logo_bytes), width=logo_width)
            except Exception:
                pass
            self._clear_table_borders(table)
            return

        # No logo: leave header empty (or role only). Name always lives in the body
        # so it never appears twice on page 1.
        if header.get("role"):
            role_para = page_header.add_paragraph()
            role_para.paragraph_format.space_before = Pt(0)
            role_para.paragraph_format.space_after = Pt(0)
            role_run = role_para.add_run(str(header["role"]))
            role_run.font.name = font_family
            role_run.font.size = Pt(body_size + 0.5)
            role_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def _fill_docx_page_footer(
        self,
        page_footer: Any,
        company_lines: list[Any],
        font_family: str,
        stamp_bytes: bytes | None = None,
    ) -> None:
        """Put company stamp/sign in the Word section footer (repeats on every page)."""
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO

        self._clear_docx_container(page_footer)
        lines = [str(line or "").strip() for line in (company_lines or []) if str(line or "").strip()]

        if stamp_bytes:
            stamp_para = page_footer.add_paragraph()
            stamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            stamp_para.paragraph_format.space_before = Pt(0)
            stamp_para.paragraph_format.space_after = Pt(2)
            stamp_run = stamp_para.add_run()
            stamp_width = Inches(1.15)
            try:
                from PIL import Image as PILImage

                pil_img = PILImage.open(BytesIO(stamp_bytes))
                img_width, img_height = pil_img.size
                aspect = img_width / img_height if img_height else 2
                if aspect > 3:
                    stamp_width = Inches(1.45)
                elif aspect < 1.1:
                    stamp_width = Inches(0.85)
            except Exception:
                pass
            stamp_run.add_picture(BytesIO(stamp_bytes), width=stamp_width)

        if not lines and not stamp_bytes:
            # Keep an empty paragraph so Word still has a valid footer part.
            page_footer.add_paragraph()
            return

        for idx, text in enumerate(lines):
            para = page_footer.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(text)
            run.font.name = font_family
            run.font.size = Pt(9 if idx > 0 else 10)
            run.bold = idx == 0
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    def _decode_logo_bytes(self, logo: dict | None, strip_black_bg: bool = False) -> bytes | None:
        import base64

        if not logo or not isinstance(logo, dict):
            return None
        logo_data = str(logo.get("data") or "")
        if not logo_data.startswith("data:image") or "," not in logo_data:
            return None
        try:
            header, encoded = logo_data.split(",", 1)
            logo_bytes = base64.b64decode(encoded)
            from PIL import Image as PILImage
            from io import BytesIO

            pil_img = PILImage.open(BytesIO(logo_bytes)).convert("RGBA")
            if strip_black_bg:
                pixels = pil_img.load()
                width, height = pil_img.size
                for y in range(height):
                    for x in range(width):
                        r, g, b, a = pixels[x, y]
                        if a > 0 and r < 35 and g < 35 and b < 35:
                            pixels[x, y] = (r, g, b, 0)
            out = BytesIO()
            pil_img.save(out, format="PNG")
            return out.getvalue()
        except Exception as exc:
            logger.warning("logo_decode_failed", error=str(exc))
            # Fallback: if payload is already a common raster format, use raw bytes.
            try:
                raw = base64.b64decode(logo_data.split(",", 1)[1])
                mime = logo_data.split(";", 1)[0].lower()
                if any(kind in mime for kind in ("png", "jpeg", "jpg", "gif")) and len(raw) >= 300:
                    return raw
            except Exception:
                pass
            return None

    def _clear_table_borders(self, table: Any) -> None:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        if tbl.tblPr is None:
            tbl.insert(0, tbl_pr)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "nil")
            borders.append(element)
        existing = tbl_pr.find(qn("w:tblBorders"))
        if existing is not None:
            tbl_pr.remove(existing)
        tbl_pr.append(borders)

    def _add_horizontal_line(self, doc: Any) -> None:
        from docx.shared import Pt

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(10)
        self._apply_bottom_border(para)

    def _apply_bottom_border(self, paragraph: Any) -> None:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        p = paragraph._p
        p_pr = p.get_or_add_pPr()
        existing = p_pr.find(qn("w:pBdr"))
        if existing is not None:
            p_pr.remove(existing)
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def _add_simple_header(
        self,
        doc: Any,
        header: dict,
        header_size: int,
        body_size: int,
        font_family: str = "Calibri",
    ) -> None:
        """Add simple header without logo."""
        from docx.shared import Pt, RGBColor

        name_para = doc.add_paragraph()
        name_para.paragraph_format.space_after = Pt(2)
        name_run = name_para.add_run(str(header.get("name") or "Candidate"))
        name_run.bold = True
        name_run.font.name = font_family
        name_run.font.size = Pt(header_size)
        name_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        if header.get("role"):
            role_para = doc.add_paragraph()
            role_para.paragraph_format.space_before = Pt(0)
            role_para.paragraph_format.space_after = Pt(0)
            role_run = role_para.add_run(str(header["role"]))
            role_run.font.name = font_family
            role_run.font.size = Pt(body_size + 0.5)
            role_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def _add_docx_section(
        self,
        doc: Any,
        section: dict[str, Any],
        header_size: float,
        body_size: float,
        font_family: str = "Calibri",
        usable_width: Any = None,
    ) -> None:
        """Add a section to the DOCX document."""
        from docx.shared import Pt, RGBColor, Inches

        section_type = section.get("type")
        title = section.get("title", "")
        content = section.get("content")

        if title:
            title_para = doc.add_paragraph()
            formatted_title = title.upper() if not title.isupper() else title
            display_title = formatted_title.rstrip(":")
            title_run = title_para.add_run(display_title)
            title_run.bold = True
            title_run.font.name = font_family
            title_run.font.size = Pt(header_size)
            title_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            title_para.paragraph_format.space_before = Pt(14)
            title_para.paragraph_format.space_after = Pt(6)
            self._apply_bottom_border(title_para)

        if not content:
            return

        if section_type == "text":
            p = doc.add_paragraph()
            run = p.add_run(str(content))
            run.font.name = font_family
            run.font.size = Pt(body_size)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            p.paragraph_format.space_after = Pt(6)

        elif section_type == "skills":
            if isinstance(content, dict):
                for category, skill_list in content.items():
                    skills = [str(item) for item in self._as_list(skill_list) if str(item).strip()]
                    if not skills:
                        continue
                    cat_para = doc.add_paragraph()
                    cat_para.paragraph_format.space_after = Pt(3)
                    cat_run = cat_para.add_run(f"{category}: ")
                    cat_run.bold = True
                    cat_run.font.name = font_family
                    cat_run.font.size = Pt(body_size)
                    cat_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    skills_run = cat_para.add_run(", ".join(skills))
                    skills_run.font.name = font_family
                    skills_run.font.size = Pt(body_size)
                    skills_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            else:
                skills = [str(item) for item in self._as_list(content) if str(item).strip()]
                if skills:
                    p = doc.add_paragraph()
                    run = p.add_run(", ".join(skills))
                    run.font.name = font_family
                    run.font.size = Pt(body_size)

        elif section_type in {"experience", "education", "projects"}:
            for item in self._as_list(content):
                if section_type == "education":
                    self._add_docx_education(doc, item, body_size, font_family, usable_width)
                else:
                    self._add_docx_record(doc, item, body_size, font_family, usable_width)

        else:
            for item in self._as_list(content):
                clean_item = self._clean_inline_text(item)
                if not clean_item:
                    continue
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(clean_item)
                run.font.name = font_family
                run.font.size = Pt(body_size)

    def _add_docx_education(
        self,
        doc: Any,
        item: Any,
        body_size: float,
        font_family: str = "Calibri",
        usable_width: Any = None,
    ) -> None:
        """Render education clearly: Degree (bold), Institution, Year right-aligned."""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_TAB_ALIGNMENT

        if not isinstance(item, dict):
            text = self._clean_inline_text(item)
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = font_family
                run.font.size = Pt(body_size)
            return

        degree = self._clean_inline_text(item.get("degree") or item.get("qualification") or "")
        institution = self._clean_inline_text(
            item.get("institution") or item.get("school") or item.get("university") or ""
        )
        year = self._clean_inline_text(
            item.get("year") or item.get("duration") or item.get("dates") or item.get("graduation_date") or ""
        )
        location = self._clean_inline_text(item.get("location") or "")
        if not degree and not institution:
            return

        # Line 1: Degree ........................ Year
        line = doc.add_paragraph()
        line.paragraph_format.space_before = Pt(6)
        line.paragraph_format.space_after = Pt(0)
        tab_pos = usable_width if usable_width is not None else Inches(6.5)
        try:
            line.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
        except Exception:
            pass
        primary = degree or institution
        run = line.add_run(primary)
        run.bold = True
        run.font.name = font_family
        run.font.size = Pt(body_size)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        if year:
            line.add_run("\t")
            date_run = line.add_run(year)
            date_run.font.name = font_family
            date_run.font.size = Pt(body_size)
            date_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Line 2: Institution, Location
        secondary_parts = [p for p in (institution if degree else "", location) if p]
        # If degree was missing and institution was primary, don't repeat institution.
        if degree and institution:
            secondary_parts = [institution] + ([location] if location and location.lower() not in institution.lower() else [])
        elif not degree and location:
            secondary_parts = [location]
        else:
            secondary_parts = []
        if secondary_parts:
            inst_line = doc.add_paragraph()
            inst_line.paragraph_format.space_before = Pt(0)
            inst_line.paragraph_format.space_after = Pt(2)
            inst_run = inst_line.add_run(", ".join(secondary_parts))
            inst_run.italic = True
            inst_run.font.name = font_family
            inst_run.font.size = Pt(body_size)
            inst_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        if item.get("cgpa"):
            gpa_line = doc.add_paragraph(style="List Bullet")
            gpa_line.paragraph_format.space_after = Pt(1)
            gpa_run = gpa_line.add_run(f"CGPA: {item['cgpa']}")
            gpa_run.font.name = font_family
            gpa_run.font.size = Pt(body_size)

    def _add_docx_record(
        self,
        doc: Any,
        item: Any,
        body_size: float,
        font_family: str = "Calibri",
        usable_width: Any = None,
    ) -> None:
        """Add an experience/education/project record with right-aligned dates."""
        from docx.shared import Pt, RGBColor, Inches, Twips
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH

        if not isinstance(item, dict):
            clean_item = self._clean_inline_text(item)
            if not clean_item:
                return
            p = doc.add_paragraph()
            run = p.add_run(clean_item)
            run.font.name = font_family
            run.font.size = Pt(body_size)
            return

        title = self._clean_inline_text(
            item.get("title")
            or item.get("role")
            or item.get("position")
            or item.get("degree")
            or item.get("name")
            or ""
        )
        company = self._clean_inline_text(
            item.get("company")
            or item.get("institution")
            or item.get("organization")
            or item.get("school")
            or item.get("university")
            or ""
        )
        duration = self._clean_inline_text(
            item.get("duration")
            or item.get("date")
            or item.get("dates")
            or item.get("year")
            or item.get("graduation_date")
            or ""
        )
        location = self._clean_inline_text(item.get("location") or "")
        project_name = self._clean_inline_text(item.get("project") or item.get("project_name") or "")

        # Company (bold) + right-aligned dates on same line.
        primary = company or title
        secondary = title if company else ""
        if primary or duration:
            line = doc.add_paragraph()
            line.paragraph_format.space_before = Pt(8)
            line.paragraph_format.space_after = Pt(0)
            tab_pos = usable_width if usable_width is not None else Inches(6.5)
            try:
                line.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
            except Exception:
                pass
            if primary:
                run = line.add_run(primary)
                run.bold = True
                run.font.name = font_family
                run.font.size = Pt(body_size)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            if duration:
                line.add_run("\t")
                date_run = line.add_run(duration)
                date_run.font.name = font_family
                date_run.font.size = Pt(body_size)
                date_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        if secondary:
            role_line = doc.add_paragraph()
            role_line.paragraph_format.space_before = Pt(0)
            role_line.paragraph_format.space_after = Pt(1)
            role_run = role_line.add_run(secondary)
            role_run.italic = True
            role_run.font.name = font_family
            role_run.font.size = Pt(body_size)
            role_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        if project_name:
            proj_line = doc.add_paragraph()
            proj_line.paragraph_format.space_before = Pt(0)
            proj_line.paragraph_format.space_after = Pt(1)
            proj_run = proj_line.add_run(project_name)
            proj_run.italic = True
            proj_run.font.name = font_family
            proj_run.font.size = Pt(body_size)

        if location:
            loc_p = doc.add_paragraph()
            loc_p.paragraph_format.space_before = Pt(0)
            loc_p.paragraph_format.space_after = Pt(2)
            loc_run = loc_p.add_run(location)
            loc_run.font.name = font_family
            loc_run.font.size = Pt(max(9.5, body_size - 0.5))
            loc_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        desc = (
            item.get("description")
            or item.get("responsibilities")
            or item.get("achievements")
            or item.get("details")
            or []
        )
        details = self._as_list(desc)
        if item.get("cgpa"):
            details.insert(0, f"CGPA: {item['cgpa']}")
        # Do not re-append Environment: tech dumps into bullets — they look dated and
        # duplicate the skills section. Technologies stay on the role object only.
        for bullet in details:
            if not bullet or not str(bullet).strip():
                continue
            # Safety: never render a paragraph stream as a single bullet.
            from app.services.structured_resume_store import expand_to_bullets

            for piece in expand_to_bullets([bullet], max_bullets=8):
                clean_bullet = self._clean_inline_text(piece)
                if not clean_bullet:
                    continue
                if re.match(r"(?i)^environment\s*:", clean_bullet):
                    continue
                bullet_p = doc.add_paragraph(style="List Bullet")
                bullet_p.paragraph_format.left_indent = Inches(0.2)
                bullet_p.paragraph_format.first_line_indent = Inches(-0.15)
                bullet_p.paragraph_format.space_after = Pt(2)
                bullet_run = bullet_p.add_run(clean_bullet)
                bullet_run.font.name = font_family
                bullet_run.font.size = Pt(body_size)
                bullet_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    def _append_section(
        self,
        story: list[Any],
        section: dict[str, Any],
        body_style: Any,
        item_title_style: Any,
        bullet_style: Any,
        paragraph_cls: Any,
        spacer_cls: Any,
    ) -> None:
        section_type = section.get("type")
        content = section.get("content")

        if section_type == "text":
            story.append(paragraph_cls(self._paragraph_text(str(content)), body_style))
            return

        if section_type == "skills":
            if isinstance(content, dict):
                for category, values in content.items():
                    skills = [str(item) for item in self._as_list(values) if str(item).strip()]
                    if skills:
                        story.append(paragraph_cls(escape(f"{category}: {', '.join(skills)}"), body_style))
            else:
                skills = [str(item) for item in self._as_list(content) if str(item).strip()]
                story.append(paragraph_cls(escape(", ".join(skills)), body_style))
            return

        if section_type in {"experience", "education", "projects"}:
            for item in self._as_list(content):
                self._append_record(story, item, item_title_style, bullet_style, paragraph_cls)
                story.append(spacer_cls(1, 3))
            return

        for item in self._as_list(content):
            story.append(paragraph_cls(f"- {escape(str(item))}", bullet_style))

    def _append_record(
        self,
        story: list[Any],
        item: Any,
        item_title_style: Any,
        bullet_style: Any,
        paragraph_cls: Any,
    ) -> None:
        if not isinstance(item, dict):
            story.append(paragraph_cls(f"- {escape(str(item))}", bullet_style))
            return

        title = (
            item.get("title")
            or item.get("role")
            or item.get("position")
            or item.get("degree")
            or item.get("name")
            or "Detail"
        )
        organization = item.get("company") or item.get("institution") or item.get("organization")
        duration = item.get("duration") or item.get("date") or item.get("year")

        title_parts = [str(title)]
        if organization:
            title_parts.append(str(organization))
        if duration:
            title_parts.append(str(duration))

        story.append(paragraph_cls(escape(" - ".join(title_parts)), item_title_style))
        description = (
            item.get("description")
            or item.get("responsibilities")
            or item.get("achievements")
            or item.get("details")
        )
        for line in self._as_list(description):
            if str(line).strip():
                story.append(paragraph_cls(f"- {escape(str(line))}", bullet_style))

    def _as_list(self, value: Any) -> list[Any]:
        if value is None:
            return []
        if isinstance(value, list):
            return value
        if isinstance(value, tuple):
            return list(value)
        if isinstance(value, dict):
            for key in ("items", "values", "details"):
                nested = value.get(key)
                if isinstance(nested, list):
                    return nested
            return [value]
        if isinstance(value, str):
            lines = [line.strip() for line in value.splitlines() if line.strip()]
            return lines or [value]
        return [value]

    def _paragraph_text(self, value: str) -> str:
        return escape(value).replace("\n", "<br/>")

    def _safe_int(self, value: Any, default: int) -> int:
        try:
            return int(float(value))
        except (TypeError, ValueError):
            return default
